from __future__ import annotations

import io
from pathlib import Path

import docx
import pytest
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from cv_generator.docx_import import load_footer, load_section
from cv_generator.errors import DocParseError
from cv_generator.models import RichBlock, RichParagraph, RichRun, RichTable
from tests.support import (
    AFTER_SECTION,
    BEFORE_HEADING,
    BODY_SIZE_PT,
    INHERITS_BOLD,
    PROJEKTLISTE_NAME,
    SAMPLE_PROJECTS,
    SPLIT_RUNS,
    SWITCHED_OFF,
    png,
    write_projektliste,
    write_styled_docx,
)


@pytest.fixture
def projektliste(tmp_path: Path) -> Path:
    return write_projektliste(tmp_path / PROJEKTLISTE_NAME)


@pytest.fixture
def blocks(projektliste: Path) -> list[RichBlock]:
    return load_section(projektliste, "Projekthistorie")


def tables(blocks: list[RichBlock]) -> list[RichTable]:
    return [block for block in blocks if isinstance(block, RichTable)]


def cell_paragraphs(blocks: list[RichBlock]) -> list[RichParagraph]:
    return [
        paragraph
        for table in tables(blocks)
        for row in table.rows
        for cell in row
        for paragraph in cell.paragraphs
    ]


def runs(blocks: list[RichBlock]) -> list[RichRun]:
    return [run for paragraph in cell_paragraphs(blocks) for run in paragraph.runs]


def texts(blocks: list[RichBlock]) -> list[str]:
    """Every paragraph's text, inside tables too, for absence assertions."""
    top_level = [block.text() for block in blocks if isinstance(block, RichParagraph)]
    return top_level + [paragraph.text() for paragraph in cell_paragraphs(blocks)]


class TestNamedEnd:
    """A build recipe may say which heading the import stops before."""

    def test_the_end_heading_is_not_imported(self, projektliste: Path) -> None:
        blocks = load_section(projektliste, "Projekthistorie", "Ausbildung")
        assert "Ausbildung" not in texts(blocks)
        assert AFTER_SECTION not in texts(blocks)
        assert len(tables(blocks)) == len(SAMPLE_PROJECTS)

    def test_it_overrides_the_formatting_rule(self, projektliste: Path) -> None:
        # Left to itself the import would stop at "Ausbildung", which is shaped
        # like the heading it started from. Naming a later paragraph carries that
        # heading in, which is how a recipe overrules the guess.
        blocks = load_section(projektliste, "Projekthistorie", AFTER_SECTION)
        assert "Ausbildung" in texts(blocks)
        assert AFTER_SECTION not in texts(blocks)

    def test_it_is_matched_case_insensitively(self, projektliste: Path) -> None:
        assert load_section(projektliste, "Projekthistorie", "AUSBILDUNG")

    def test_it_may_be_a_regular_expression(self, projektliste: Path) -> None:
        blocks = load_section(projektliste, "Projekthistorie", "Ausbild.*")
        assert "Ausbildung" not in texts(blocks)

    def test_a_regular_expression_still_matches_the_whole_heading(self, projektliste: Path) -> None:
        # Not a substring search: "Ausbild" alone must not match "Ausbildung" --
        # otherwise a paragraph merely mentioning a heading's name would be
        # mistaken for the heading itself.
        with pytest.raises(DocParseError, match="no heading matching 'Ausbild' follows"):
            load_section(projektliste, "Projekthistorie", "Ausbild")

    def test_an_invalid_regular_expression_is_a_clean_error(self, projektliste: Path) -> None:
        with pytest.raises(DocParseError, match="not a valid regular expression"):
            load_section(projektliste, "Projekthistorie", "Ausbildung(")

    def test_an_end_that_never_comes_is_an_error(self, projektliste: Path) -> None:
        # Not "run to the end of the document": the recipe said where to stop,
        # so importing the rest would put another section's content in this one.
        with pytest.raises(DocParseError, match="no heading matching 'Publikationen' follows"):
            load_section(projektliste, "Projekthistorie", "Publikationen")

    def test_an_end_before_the_beginning_does_not_count(self, projektliste: Path) -> None:
        with pytest.raises(DocParseError, match="no heading matching 'Profil' follows"):
            load_section(projektliste, "Projekthistorie", "Profil")

    def test_end_may_match_a_line_whose_wording_varies(self, tmp_path: Path) -> None:
        # A form letter's closing line names a value (here, a payment term) that
        # changes from one letter to the next, so no fixed string names it -- a
        # wildcard standing in for the variable part still covers the whole line.
        document = docx.Document()
        document.add_paragraph("Vielen Dank für Ihr Vertrauen.")
        document.add_paragraph("Bitte überweisen Sie den Betrag innerhalb von 14 Tagen.")
        document.add_paragraph("Anlage: Rechnungsdetails")
        path = tmp_path / "brief.docx"
        document.save(str(path))

        blocks = load_section(
            path, end=r"Bitte überweisen Sie den Betrag innerhalb von \d+ Tagen\."
        )
        assert "Vielen Dank für Ihr Vertrauen." in texts(blocks)
        assert "Bitte überweisen" not in texts(blocks)
        assert "Anlage: Rechnungsdetails" not in texts(blocks)


class TestSectionBoundaries:
    def test_one_table_per_project(self, blocks: list[RichBlock]) -> None:
        assert len(tables(blocks)) == len(SAMPLE_PROJECTS)

    def test_the_blank_paragraphs_between_tables_are_kept(self, blocks: list[RichBlock]) -> None:
        # The source needs one to keep two tables from merging when Word
        # renders them, and it is otherwise indistinguishable from any other
        # blank line -- so, like any other blank paragraph, it is kept.
        assert len(blocks) == 2 * len(SAMPLE_PROJECTS) + 1

    def test_earlier_sections_are_not_imported(self, blocks: list[RichBlock]) -> None:
        assert BEFORE_HEADING not in texts(blocks)

    def test_import_stops_at_the_next_heading(self, blocks: list[RichBlock]) -> None:
        assert AFTER_SECTION not in texts(blocks)

    def test_heading_is_matched_case_insensitively(self, projektliste: Path) -> None:
        assert load_section(projektliste, "projekthistorie")

    def test_heading_may_be_a_regular_expression(self, projektliste: Path) -> None:
        assert load_section(projektliste, "Projekt.*orie")

    def test_blank_paragraph_inside_a_cell_is_kept(self, blocks: list[RichBlock]) -> None:
        period = tables(blocks)[0].rows[0][0]
        assert [paragraph.text() for paragraph in period.paragraphs] == [
            "02/2026 – 07/2026",
            "Rolle:",
            "programmweiter Testmanager",
            "",
            "Kunde:",
            "Land Schleswig-Holstein",
        ]

    def test_missing_heading_is_an_error(self, projektliste: Path) -> None:
        with pytest.raises(DocParseError, match="no heading matching 'Projekte' found"):
            load_section(projektliste, "Projekte")

    def test_a_section_of_only_blank_paragraphs_is_not_empty(self, tmp_path: Path) -> None:
        # Blank paragraphs are content now (they preserve the source's own
        # spacing), so a section that is only blank lines is not the same as
        # no section at all.
        path = write_projektliste(tmp_path / PROJEKTLISTE_NAME, [])
        blocks = load_section(path, "Projekthistorie", "Ausbildung")
        assert len(blocks) == 1
        assert blocks[0].text() == ""  # type: ignore[union-attr]

    def test_a_heading_immediately_followed_by_the_end_is_an_error(self, tmp_path: Path) -> None:
        # Nothing at all between the two headings, not even a blank paragraph.
        document = docx.Document()
        document.add_paragraph("Start")
        document.add_paragraph("End")
        path = tmp_path / "brief.docx"
        document.save(str(path))
        with pytest.raises(DocParseError, match="nothing follows the heading"):
            load_section(path, "Start", "End")

    def test_a_file_that_is_not_a_word_document(self, tmp_path: Path) -> None:
        path = tmp_path / "Projektliste.docx"
        path.write_text("just text", encoding="utf-8")
        with pytest.raises(DocParseError, match="as a Word document"):
            load_section(path, "Projekthistorie")


class TestFormatting:
    def test_bold_and_size_are_carried_over(self, blocks: list[RichBlock]) -> None:
        period = tables(blocks)[0].rows[0][0].paragraphs
        assert [(run.text, run.bold, run.size_pt) for run in period[0].runs] == [
            ("02/2026 – 07/2026", True, BODY_SIZE_PT)
        ]
        assert period[2].runs[0].bold is False

    def test_style_is_resolved_and_a_run_can_switch_it_off(self, tmp_path: Path) -> None:
        path = write_styled_docx(tmp_path / "Projektliste.docx")
        styled = load_section(path, "Projekthistorie")[0]
        assert isinstance(styled, RichParagraph)
        assert [(run.text, run.bold) for run in styled.runs] == [
            (INHERITS_BOLD, True),
            (SWITCHED_OFF, False),
        ]

    def test_neighbouring_runs_of_equal_formatting_are_merged(self, tmp_path: Path) -> None:
        # Word splits a sentence at every spell-check marker and revision; left
        # alone, each fragment would become its own <span> and its own Word run.
        path = write_styled_docx(tmp_path / "Projektliste.docx")
        split = load_section(path, "Projekthistorie")[1]
        assert isinstance(split, RichParagraph)
        assert [run.text for run in split.runs] == ["".join(SPLIT_RUNS)]

    def test_hyperlink_keeps_its_text_and_target(self, blocks: list[RichBlock]) -> None:
        linked = [run for run in runs(blocks) if run.link]
        assert [(run.text, run.link) for run in linked] == [
            ("Projektseite", "https://example.org/projekt")
        ]


class TestLists:
    def test_bullets_become_list_items_at_their_level(self, blocks: list[RichBlock]) -> None:
        detail = tables(blocks)[0].rows[0][1].paragraphs
        assert [(p.text(), p.level, p.ordered) for p in detail] == [
            ("Tätigkeiten", None, False),
            ("Konzeption einer Integrationstestumgebung", 0, False),
            ("Aufbau der Testautomatisierung", 0, False),
            ("mit Playwright & TypeScript", 1, False),
            ("Projektseite", None, False),
        ]


class TestTableGeometry:
    def test_two_columns_with_their_widths(self, blocks: list[RichBlock]) -> None:
        table = tables(blocks)[0]
        assert [len(row) for row in table.rows] == [2]
        assert table.column_widths == pytest.approx([0.2, 0.8], abs=0.005)

    def test_a_ruled_table_says_so(self, blocks: list[RichBlock]) -> None:
        # `Table Grid` keeps its borders in the style, not on the table.
        assert all(table.bordered for table in tables(blocks))

    def test_a_docx_table_is_not_centered(self, blocks: list[RichBlock]) -> None:
        # Unlike an `.xlsx` range, a `.docx` table is meant to fill the page the
        # way it did in the source document.
        assert all(not table.centered for table in tables(blocks))


def _write_footer_docx(path: Path, *, footer_paragraphs: list[str] | None = None) -> None:
    """A minimal `.docx` with a body paragraph and, optionally, a footer."""
    document = docx.Document()
    document.add_paragraph("Body content.")
    if footer_paragraphs:
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False
        footer.paragraphs[0].text = footer_paragraphs[0]
        for text in footer_paragraphs[1:]:
            footer.add_paragraph(text)
    document.save(str(path))


class TestFooter:
    def test_a_plain_footer_paragraph_is_read(self, tmp_path: Path) -> None:
        path = tmp_path / "brief.docx"
        _write_footer_docx(path, footer_paragraphs=["Footer line"])
        blocks = load_footer(path)
        assert len(blocks) == 1
        assert blocks[0].text() == "Footer line"  # type: ignore[union-attr]

    def test_a_document_with_no_footer_definition_returns_nothing(self, tmp_path: Path) -> None:
        path = tmp_path / "brief.docx"
        _write_footer_docx(path)
        assert load_footer(path) == []

    def test_blank_footer_paragraphs_are_kept(self, tmp_path: Path) -> None:
        path = tmp_path / "brief.docx"
        _write_footer_docx(path, footer_paragraphs=["", "Real line", ""])
        blocks = load_footer(path)
        assert [b.text() for b in blocks] == ["", "Real line", ""]  # type: ignore[union-attr]

    def test_formatting_is_preserved(self, tmp_path: Path) -> None:
        document = docx.Document()
        document.add_paragraph("Body")
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False
        run = footer.paragraphs[0].add_run("Bold footer")
        run.bold = True
        path = tmp_path / "brief.docx"
        document.save(str(path))
        blocks = load_footer(path)
        assert blocks[0].runs[0].bold is True  # type: ignore[union-attr]

    def test_the_first_page_footer_wins_when_it_has_content(self, tmp_path: Path) -> None:
        # A single-page letter shows only its first-page footer; the "default"
        # one, meant for a second page that never comes, is typically empty.
        document = docx.Document()
        document.add_paragraph("Body")
        section = document.sections[0]
        section.different_first_page_header_footer = True
        section.first_page_footer.is_linked_to_previous = False
        section.first_page_footer.paragraphs[0].text = "First-page footer"
        path = tmp_path / "brief.docx"
        document.save(str(path))
        blocks = load_footer(path)
        assert len(blocks) == 1
        assert blocks[0].text() == "First-page footer"  # type: ignore[union-attr]

    def test_the_default_footer_is_used_when_the_first_page_one_is_empty(
        self, tmp_path: Path
    ) -> None:
        document = docx.Document()
        document.add_paragraph("Body")
        section = document.sections[0]
        section.different_first_page_header_footer = True
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = "Default footer"
        path = tmp_path / "brief.docx"
        document.save(str(path))
        blocks = load_footer(path)
        assert len(blocks) == 1
        assert blocks[0].text() == "Default footer"  # type: ignore[union-attr]

    def test_a_text_box_s_alternate_content_is_read_only_once(self, tmp_path: Path) -> None:
        # A letterhead-style footer commonly anchors its content in a text box,
        # written twice in the XML -- once as the modern shape (`mc:Choice`),
        # once as a legacy fallback (`mc:Fallback`) for a Word too old to read
        # the first. Only one is ever shown, and only one may be read.
        document = docx.Document()
        document.add_paragraph("Body")
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False
        run_xml = f"""
        <w:r {nsdecls("w")}
             xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">
          <mc:AlternateContent>
            <mc:Choice Requires="wps">
              <w:pict><w:txbxContent>
                <w:p><w:r><w:t>Choice text</w:t></w:r></w:p>
              </w:txbxContent></w:pict>
            </mc:Choice>
            <mc:Fallback>
              <w:pict><w:txbxContent>
                <w:p><w:r><w:t>Fallback text</w:t></w:r></w:p>
              </w:txbxContent></w:pict>
            </mc:Fallback>
          </mc:AlternateContent>
        </w:r>
        """
        footer.paragraphs[0]._p.append(parse_xml(run_xml))
        path = tmp_path / "brief.docx"
        document.save(str(path))
        texts = [block.text() for block in load_footer(path)]  # type: ignore[union-attr]
        assert "Choice text" in texts
        assert "Fallback text" not in texts


def _top_level_runs(blocks: list[RichBlock]) -> list[RichRun]:
    return [run for block in blocks if isinstance(block, RichParagraph) for run in block.runs]


class TestImages:
    """A picture in a source document -- typically a photo in an identity
    table cell, as in the real ``cv-head.docx`` -- is read into the run that
    carries it, not silently dropped the way a text-only run reader would drop
    it (a picture-only run's ``text`` is empty). ``add_picture`` writes an
    inline picture (``wp:inline``); a floating one (``wp:anchor``, what
    ``cv-head.docx`` actually uses) wraps the same ``pic:pic`` differently, but
    ``_run_image`` finds the ``a:blip`` either way without caring which."""

    def test_an_inline_image_is_read(self, tmp_path: Path) -> None:
        document = docx.Document()
        run = document.add_paragraph().add_run()
        run.add_picture(io.BytesIO(png(20, 20)))
        path = tmp_path / "brief.docx"
        document.save(str(path))

        blocks = load_section(path)
        image_runs = [r for r in _top_level_runs(blocks) if r.image is not None]
        assert len(image_runs) == 1
        assert image_runs[0].image is not None
        assert image_runs[0].image.media_type == "image/png"

    def test_an_image_run_with_no_text_is_not_dropped(self, tmp_path: Path) -> None:
        # A picture-only run has empty `.text`; only a text-blind filter would
        # mistake that for an empty, droppable run.
        document = docx.Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Before. ")
        picture_run = paragraph.add_run()
        picture_run.add_picture(io.BytesIO(png(20, 20)))
        paragraph.add_run(" After.")
        path = tmp_path / "brief.docx"
        document.save(str(path))

        blocks = load_section(path)
        found = [r for r in _top_level_runs(blocks) if r.image is not None]
        assert len(found) == 1

    def test_image_bytes_round_trip(self, tmp_path: Path) -> None:
        data = png(20, 20)
        document = docx.Document()
        run = document.add_paragraph().add_run()
        run.add_picture(io.BytesIO(data))
        path = tmp_path / "brief.docx"
        document.save(str(path))

        blocks = load_section(path)
        image_runs = [r for r in _top_level_runs(blocks) if r.image is not None]
        assert image_runs[0].image is not None
        assert image_runs[0].image.data == data

    def test_an_image_in_a_table_cell_is_read(self, tmp_path: Path) -> None:
        # The real layout: identity text in one cell, a photo in the next.
        document = docx.Document()
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].paragraphs[0].add_run("Ada Lovelace")
        run = table.rows[0].cells[1].paragraphs[0].add_run()
        run.add_picture(io.BytesIO(png(20, 20)))
        path = tmp_path / "brief.docx"
        document.save(str(path))

        blocks = load_section(path)
        assert any(r.image is not None for r in runs(blocks))

    def test_an_image_run_is_not_merged_with_a_neighbouring_text_run(self, tmp_path: Path) -> None:
        # Merging equally-formatted neighbours (see TestFormatting) must not
        # concatenate a picture run's (empty) text onto a plain one and lose
        # the picture in the process.
        document = docx.Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("Before")
        picture_run = paragraph.add_run()
        picture_run.add_picture(io.BytesIO(png(20, 20)))
        path = tmp_path / "brief.docx"
        document.save(str(path))

        blocks = load_section(path)
        paragraph_runs = _top_level_runs(blocks)
        assert any(r.image is not None for r in paragraph_runs)
        assert any(r.text == "Before" for r in paragraph_runs)
