"""The build recipe: reading `config.json`, and assembling the document it describes."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pytest

from cv_generator.config import BuildConfig, SectionSpec, load_config, resolve_source
from cv_generator.errors import DocParseError
from cv_generator.models import Document, RichTable
from cv_generator.parser import build_doc, load_doc, slugify
from tests.conftest import PROJECTS_CONFIG, PROJECTS_MD, write_config
from tests.support import (
    AFTER_SECTION,
    BEFORE_HEADING,
    PROJEKTLISTE_NAME,
    SAMPLE_PROJECTS,
    png,
    write_projektliste,
    write_workbook,
)

# A Markdown source with four headings, so a span can start late, stop early, or
# cover several at once.
FOUR_SECTIONS = """---
name: Ada Lovelace
---

Ein Kurzprofil.

## Berufserfahrung

- Analyst

## Projekte

Wird von keinem Eintrag angefordert.

## Kenntnisse

- Python

## Sprachen

- Deutsch
"""

# An entry that takes document.md's header and none of its sections, by ending at the
# very first heading. Prepended by `recipe` so a test about spans can say only
# what it is about; the header has to come from somewhere in every recipe.
HEADER_ONLY: dict[str, Any] = {"source": "document.md", "format": "md", "end": "Berufserfahrung"}


def recipe(*specs: dict[str, Any]) -> dict[str, Any]:
    return {"sections": [HEADER_ONLY, *specs]}


def build(directory: Path, config: dict[str, Any]) -> Document:
    return load_doc(write_config(directory, config)).doc


def spans(directory: Path, *specs: dict[str, Any]) -> Document:
    return build(directory, recipe(*specs))


@pytest.fixture
def four_sections(tmp_path: Path) -> Path:
    (tmp_path / "document.md").write_text(FOUR_SECTIONS, encoding="utf-8")
    return tmp_path


class TestLoadConfig:
    def test_reads_the_recipe(self, tmp_path: Path) -> None:
        config = load_config(write_config(tmp_path, PROJECTS_CONFIG))
        assert [spec.source for spec in config.sections] == [
            "document.md",
            PROJEKTLISTE_NAME,
            "document.md",
        ]

    def test_begin_end_and_title_are_all_optional(self, tmp_path: Path) -> None:
        config = load_config(
            write_config(tmp_path, {"sections": [{"source": "document.md", "format": "md"}]})
        )
        assert config.sections[0].begin is None
        assert config.sections[0].end is None
        assert config.sections[0].title is None

    def test_source_and_format_are_the_required_keys_of_an_entry(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="source"):
            load_config(write_config(tmp_path, {"sections": [{"format": "md", "begin": "A"}]}))

    def test_format_is_required(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="format"):
            load_config(write_config(tmp_path, {"sections": [{"source": "document.md"}]}))

    def test_rejects_an_unknown_key(self, tmp_path: Path) -> None:
        # A typo in a recipe must fail loudly, the way frontmatter does; silently
        # dropping "sektions" would render a document with no sections at all.
        path = write_config(tmp_path, {"sections": [], "outputs": "doc"})
        with pytest.raises(DocParseError, match="outputs"):
            load_config(path)

    def test_metadata_is_no_longer_a_key(self, tmp_path: Path) -> None:
        # It was replaced by an entry with no `begin`, and leaving both in would
        # be two ways to say where the header comes from.
        path = write_config(
            tmp_path,
            {"metadata": "document.md", "sections": [{"source": "document.md", "format": "md"}]},
        )
        with pytest.raises(DocParseError, match="metadata"):
            load_config(path)

    def test_rejects_an_unknown_key_in_a_section(self, tmp_path: Path) -> None:
        config = {"sections": [{"source": "document.md", "format": "md", "from": "A"}]}
        with pytest.raises(DocParseError, match="from"):
            load_config(write_config(tmp_path, config))

    def test_rejects_a_recipe_without_sections(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="sections"):
            load_config(write_config(tmp_path, {"sections": []}))

    def test_rejects_invalid_json(self, tmp_path: Path) -> None:
        path = tmp_path / "config.json"
        path.write_text('{"sections": [],}', encoding="utf-8")
        with pytest.raises(DocParseError, match="invalid JSON"):
            load_config(path)

    def test_rejects_a_json_document_that_is_not_an_object(self, tmp_path: Path) -> None:
        path = tmp_path / "config.json"
        path.write_text("[1, 2]", encoding="utf-8")
        with pytest.raises(DocParseError, match="must be a JSON object"):
            load_config(path)

    def test_missing_file_names_the_path(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="cannot read"):
            load_config(tmp_path / "nope.json")

    def test_a_recipe_needs_nothing_but_sections(self) -> None:
        config = BuildConfig(sections=[SectionSpec(source="document.md", format="md")])
        assert config.output is None
        assert config.target is None

    def test_rejects_output_and_target_together(self, tmp_path: Path) -> None:
        # Both answer "what are the results called"; target also says where, so
        # giving both would silently ignore one of them.
        config = {
            "sections": [{"source": "document.md", "format": "md"}],
            "output": "doc",
            "target": "exports/doc",
        }
        with pytest.raises(DocParseError, match="not both"):
            load_config(write_config(tmp_path, config))

    def test_rejects_noframes_on_a_markdown_entry(self, tmp_path: Path) -> None:
        config = {"sections": [{"source": "document.md", "format": "md", "noframes": True}]}
        with pytest.raises(DocParseError, match="'noframes' only applies to format 'xlsx'"):
            load_config(write_config(tmp_path, config))

    def test_rejects_noframes_on_a_docx_entry(self, tmp_path: Path) -> None:
        config = {
            "sections": [{"source": "Projektliste.docx", "format": "docx", "noframes": False}]
        }
        with pytest.raises(DocParseError, match="'noframes' only applies to format 'xlsx'"):
            load_config(write_config(tmp_path, config))

    def test_accepts_noframes_on_an_xlsx_entry(self, tmp_path: Path) -> None:
        config = {
            "sections": [
                {"source": "Rechnung.xlsx", "format": "xlsx", "noframes": True, **XLSX_RANGE}
            ]
        }
        loaded = load_config(write_config(tmp_path, config))
        assert loaded.sections[0].noframes is True

    def test_noframes_is_optional_and_defaults_to_unset(self, tmp_path: Path) -> None:
        config = {"sections": [{"source": "Rechnung.xlsx", "format": "xlsx", **XLSX_RANGE}]}
        loaded = load_config(write_config(tmp_path, config))
        assert loaded.sections[0].noframes is None

    @pytest.mark.parametrize("spelling", ["true", "True", "TRUE", "false", "False", "FALSE"])
    def test_noframes_accepts_case_insensitive_string_spellings(
        self, tmp_path: Path, spelling: str
    ) -> None:
        config = {
            "sections": [
                {"source": "Rechnung.xlsx", "format": "xlsx", "noframes": spelling, **XLSX_RANGE}
            ]
        }
        loaded = load_config(write_config(tmp_path, config))
        assert loaded.sections[0].noframes is (spelling.lower() == "true")


class TestResolveSource:
    def test_a_plain_name_is_a_path(self, tmp_path: Path) -> None:
        (tmp_path / "document.md").write_text("x", encoding="utf-8")
        assert resolve_source(tmp_path, "document.md") == tmp_path / "document.md"

    def test_a_subdirectory_works(self, tmp_path: Path) -> None:
        (tmp_path / "sub").mkdir()
        (tmp_path / "sub" / "document.md").write_text("x", encoding="utf-8")
        assert resolve_source(tmp_path, "sub/document.md") == tmp_path / "sub" / "document.md"

    def test_an_absolute_path_is_used_as_is(self, tmp_path: Path) -> None:
        target = tmp_path / "document.md"
        target.write_text("x", encoding="utf-8")
        assert resolve_source(Path("/nowhere"), str(target)) == target

    def test_a_missing_file_names_the_path(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="no such file"):
            resolve_source(tmp_path, "document.md")

    def test_a_glob_finds_the_one_match(self, tmp_path: Path) -> None:
        # The point of allowing one: the real project list carries a date in its
        # name, so a literal name would go stale every time it is reissued.
        (tmp_path / "Bergemann-Projektliste_19_08_2026.docx").write_bytes(b"")
        found = resolve_source(tmp_path, "*Projektliste*.docx")
        assert found.name == "Bergemann-Projektliste_19_08_2026.docx"

    def test_a_glob_matching_nothing_is_an_error(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="matches"):
            resolve_source(tmp_path, "*Projektliste*.docx")

    def test_a_glob_matching_several_is_an_error(self, tmp_path: Path) -> None:
        # Picking one would quietly publish a Document built from last year's list.
        (tmp_path / "Projektliste_alt.docx").write_bytes(b"")
        (tmp_path / "Projektliste_neu.docx").write_bytes(b"")
        with pytest.raises(DocParseError, match="keep exactly one"):
            resolve_source(tmp_path, "*Projektliste*.docx")

    def test_word_s_lock_file_does_not_count_as_a_match(self, tmp_path: Path) -> None:
        # Otherwise having the document open in Word would break every build --
        # exactly when someone is most likely to rebuild.
        (tmp_path / "Projektliste.docx").write_bytes(b"")
        (tmp_path / "~$ojektliste.docx").write_bytes(b"")
        assert resolve_source(tmp_path, "*rojektliste*.docx").name == "Projektliste.docx"

    def test_an_empty_reference_is_an_error(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match="empty string"):
            resolve_source(tmp_path, "  ")

    def test_error_names_the_entry_it_came_from(self, tmp_path: Path) -> None:
        with pytest.raises(DocParseError, match=r"sections\[2\]"):
            resolve_source(tmp_path, "gone.md", source="config.json: sections[2]")

    def test_project_root_is_tried_when_base_dir_misses(self, tmp_path: Path) -> None:
        base_dir = tmp_path / "recipe"
        base_dir.mkdir()
        root = tmp_path / "root"
        root.mkdir()
        (root / "document.md").write_text("x", encoding="utf-8")
        found = resolve_source(base_dir, "document.md", project_root=root)
        assert found == root / "document.md"

    def test_base_dir_wins_over_project_root(self, tmp_path: Path) -> None:
        base_dir = tmp_path / "recipe"
        base_dir.mkdir()
        (base_dir / "document.md").write_text("near", encoding="utf-8")
        root = tmp_path / "root"
        root.mkdir()
        (root / "document.md").write_text("far", encoding="utf-8")
        found = resolve_source(base_dir, "document.md", project_root=root)
        assert found == base_dir / "document.md"

    def test_missing_everywhere_still_names_the_base_dir_path(self, tmp_path: Path) -> None:
        base_dir = tmp_path / "recipe"
        base_dir.mkdir()
        root = tmp_path / "root"
        root.mkdir()
        with pytest.raises(DocParseError, match="no such file"):
            resolve_source(base_dir, "gone.md", project_root=root)

    def test_project_root_fallback_does_not_apply_to_a_glob(self, tmp_path: Path) -> None:
        base_dir = tmp_path / "recipe"
        base_dir.mkdir()
        root = tmp_path / "root"
        root.mkdir()
        (root / "Projektliste.docx").write_bytes(b"")
        with pytest.raises(DocParseError, match="matches"):
            resolve_source(base_dir, "*Projektliste*.docx", project_root=root)


class TestMarkdownSpans:
    """`begin` and `end` pick a span out of a Markdown source."""

    def test_a_span_stops_before_its_end_headline(self, four_sections: Path) -> None:
        doc = spans(
            four_sections,
            {
                "source": "document.md",
                "format": "md",
                "begin": "Berufserfahrung",
                "end": "Projekte",
            },
        )
        assert [s.title for s in doc.sections] == ["Berufserfahrung"]

    def test_a_span_without_an_end_runs_to_the_end_of_the_file(self, four_sections: Path) -> None:
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "Kenntnisse"})
        assert [s.title for s in doc.sections] == ["Kenntnisse", "Sprachen"]

    def test_a_span_without_a_begin_starts_at_the_top_of_the_file(
        self, four_sections: Path
    ) -> None:
        doc = build(
            four_sections,
            {"sections": [{"source": "document.md", "format": "md", "end": "Kenntnisse"}]},
        )
        assert [s.title for s in doc.sections] == ["Berufserfahrung", "Projekte"]

    def test_a_span_with_neither_is_the_whole_file(self, four_sections: Path) -> None:
        doc = build(four_sections, {"sections": [{"source": "document.md", "format": "md"}]})
        assert [s.title for s in doc.sections] == [
            "Berufserfahrung",
            "Projekte",
            "Kenntnisse",
            "Sprachen",
        ]

    def test_ending_at_the_first_heading_takes_the_header_and_no_sections(
        self, four_sections: Path
    ) -> None:
        # The idiom that replaces the old `metadata` key: an entry that wants a
        # file's frontmatter but none of its content.
        doc = build(
            four_sections,
            {"sections": [{"source": "document.md", "format": "md", "end": "Berufserfahrung"}]},
        )
        assert doc.name == "Ada Lovelace"
        assert doc.sections == []

    def test_a_span_is_split_at_its_own_headings(self, four_sections: Path) -> None:
        # One entry, several sections: each keeps its own heading and anchor
        # rather than becoming body text of the first.
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "Kenntnisse"})
        assert [s.slug for s in doc.sections] == ["kenntnisse", "sprachen"]
        assert "## Sprachen" not in (doc.sections[0].markdown)

    def test_content_no_entry_asks_for_is_left_out(self, four_sections: Path) -> None:
        doc = spans(
            four_sections,
            {
                "source": "document.md",
                "format": "md",
                "begin": "Berufserfahrung",
                "end": "Projekte",
            },
            {"source": "document.md", "format": "md", "begin": "Kenntnisse"},
        )
        assert all("keinem Eintrag" not in section.markdown for section in doc.sections)

    def test_the_same_file_may_be_used_more_than_once(self, four_sections: Path) -> None:
        doc = spans(
            four_sections,
            {"source": "document.md", "format": "md", "begin": "Sprachen"},
            {
                "source": "document.md",
                "format": "md",
                "begin": "Berufserfahrung",
                "end": "Projekte",
            },
        )
        assert [s.title for s in doc.sections] == ["Sprachen", "Berufserfahrung"]

    def test_a_headline_may_be_written_with_its_hashes(self, four_sections: Path) -> None:
        doc = spans(
            four_sections, {"source": "document.md", "format": "md", "begin": "## Sprachen"}
        )
        assert [s.title for s in doc.sections] == ["Sprachen"]

    def test_matching_ignores_case(self, four_sections: Path) -> None:
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "sprachen"})
        assert [s.title for s in doc.sections] == ["Sprachen"]

    def test_begin_may_be_a_regular_expression(self, four_sections: Path) -> None:
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "Sprach.n"})
        assert [s.title for s in doc.sections] == ["Sprachen"]

    def test_end_may_be_a_regular_expression(self, four_sections: Path) -> None:
        doc = spans(
            four_sections,
            {"source": "document.md", "format": "md", "begin": "Berufserfahrung", "end": "Proj.*"},
        )
        assert [s.title for s in doc.sections] == ["Berufserfahrung"]

    def test_a_regular_expression_still_matches_the_whole_headline(
        self, four_sections: Path
    ) -> None:
        # Not a substring search: "Sprach" alone must not match "Sprachen", the
        # same as a plain, misspelled headline never has before -- otherwise
        # "Projekte" would also match a later "Weitere Projekte" heading.
        with pytest.raises(DocParseError, match="no heading matching 'Sprach'"):
            spans(four_sections, {"source": "document.md", "format": "md", "begin": "Sprach"})

    def test_an_invalid_regular_expression_is_a_clean_error(self, four_sections: Path) -> None:
        with pytest.raises(DocParseError, match="not a valid regular expression"):
            spans(four_sections, {"source": "document.md", "format": "md", "begin": "Sprachen("})

    def test_the_section_keeps_the_source_heading_s_own_wording(self, four_sections: Path) -> None:
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "sprachen"})
        assert doc.sections[0].title == "Sprachen"

    def test_title_renames_the_first_section_of_a_span(self, four_sections: Path) -> None:
        doc = spans(
            four_sections,
            {"source": "document.md", "format": "md", "begin": "Kenntnisse", "title": "Skills"},
        )
        assert [s.title for s in doc.sections] == ["Skills", "Sprachen"]
        assert doc.sections[0].slug == "skills"

    def test_a_missing_begin_headline_lists_the_ones_there_are(self, four_sections: Path) -> None:
        with pytest.raises(
            DocParseError, match="no heading matching 'Hobbys'; it has 'Berufserfahrung'"
        ):
            spans(four_sections, {"source": "document.md", "format": "md", "begin": "Hobbys"})

    def test_an_end_headline_that_never_comes_is_an_error(self, four_sections: Path) -> None:
        # Not "run to the end of the file": the entry said where to stop, and
        # importing the rest would put another section's content in this one.
        with pytest.raises(
            DocParseError, match="no heading matching 'Hobbys' after heading matching 'Kenntnisse'"
        ):
            spans(
                four_sections,
                {"source": "document.md", "format": "md", "begin": "Kenntnisse", "end": "Hobbys"},
            )

    def test_an_end_headline_before_begin_does_not_count(self, four_sections: Path) -> None:
        with pytest.raises(DocParseError, match="no heading matching 'Projekte' after"):
            spans(
                four_sections,
                {"source": "document.md", "format": "md", "begin": "Kenntnisse", "end": "Projekte"},
            )

    def test_a_missing_end_names_the_top_of_the_file_when_there_is_no_begin(
        self, four_sections: Path
    ) -> None:
        with pytest.raises(
            DocParseError, match="no heading matching 'Hobbys' after the top of the file"
        ):
            build(
                four_sections,
                {"sections": [{"source": "document.md", "format": "md", "end": "Hobbys"}]},
            )

    def test_a_source_needs_no_frontmatter_of_its_own(self, tmp_path: Path) -> None:
        # Only the file supplying the header is a Document; anything else is Markdown.
        (tmp_path / "document.md").write_text(FOUR_SECTIONS, encoding="utf-8")
        (tmp_path / "extra.md").write_text("## Hobbys\n\n- Segeln\n", encoding="utf-8")
        doc = spans(tmp_path, {"source": "extra.md", "format": "md", "begin": "Hobbys"})
        assert doc.sections[0].markdown == "- Segeln"

    def test_the_file_each_section_came_from_is_recorded(self, four_sections: Path) -> None:
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "Sprachen"})
        assert doc.sections[0].source == str(four_sections / "document.md")

    def test_duplicate_titles_from_two_files_get_distinct_slugs(self, tmp_path: Path) -> None:
        # Slugs are the HTML anchors, and two files easily use the same heading.
        (tmp_path / "document.md").write_text(FOUR_SECTIONS, encoding="utf-8")
        (tmp_path / "other.md").write_text("## Kenntnisse\n\n- Rust\n", encoding="utf-8")
        doc = spans(
            tmp_path,
            {"source": "document.md", "format": "md", "begin": "Kenntnisse", "end": "Sprachen"},
            {"source": "other.md", "format": "md", "begin": "Kenntnisse"},
        )
        assert [s.slug for s in doc.sections] == ["kenntnisse", "kenntnisse-2"]


class TestWordSpans:
    """A `.docx` entry brings in blocks, not Markdown."""

    def test_the_section_carries_the_imported_blocks(self, projects_document: Document) -> None:
        # Blank paragraphs from the source are interspersed and kept (see
        # TestSectionBoundaries in test_docx_import.py), so this counts tables only.
        projects = projects_document.section("projekte")
        assert projects is not None
        tables = [block for block in projects.blocks if isinstance(block, RichTable)]
        assert len(tables) == len(SAMPLE_PROJECTS)

    def test_the_section_has_no_markdown(self, projects_document: Document) -> None:
        # Two sources for one section is how a stale paragraph ends up in
        # whichever output format happens to prefer one over the other.
        projects = projects_document.section("projekte")
        assert projects is not None
        assert projects.markdown == ""

    def test_the_source_file_is_recorded(self, projects_document: Document) -> None:
        projects = projects_document.section("projekte")
        assert projects is not None
        assert projects.source is not None
        assert projects.source.endswith(PROJEKTLISTE_NAME)

    def test_title_renames_the_imported_section(self, projects_document: Document) -> None:
        # "Projekthistorie" in Word, "Projekte" in the Document: without the override
        # the target would have to adopt the source document's wording.
        titles = [s.title for s in projects_document.sections]
        assert titles == ["Berufserfahrung", "Projekte", "Kenntnisse"]

    def test_without_a_title_the_begin_headline_is_used(self, projects_dir: Path) -> None:
        doc = spans(
            projects_dir,
            {
                "source": PROJEKTLISTE_NAME,
                "format": "docx",
                "begin": "Projekthistorie",
                "end": "Ausbildung",
            },
        )
        assert [s.title for s in doc.sections] == ["Projekthistorie"]

    def test_the_end_headline_bounds_the_import(self, projects_dir: Path) -> None:
        # "Ausbildung" is the heading after the project list; its one paragraph
        # must not be dragged in with the projects.
        doc = spans(
            projects_dir,
            {
                "source": PROJEKTLISTE_NAME,
                "format": "docx",
                "begin": "Projekthistorie",
                "end": "Ausbildung",
            },
        )
        assert "Promotion Maschinenbau" not in json.dumps(doc.sections[0].model_dump(mode="json"))

    def test_omitting_the_end_falls_back_to_the_formatting_rule(self, projects_dir: Path) -> None:
        # A hand-made Word Document has no heading styles, so the import stops at the
        # next paragraph shaped like the heading it started from.
        doc = spans(
            projects_dir,
            {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"},
        )
        tables = [block for block in doc.sections[0].blocks if isinstance(block, RichTable)]
        assert len(tables) == len(SAMPLE_PROJECTS)

    def test_omitting_the_begin_starts_at_the_top_of_the_document(self, projects_dir: Path) -> None:
        # No heading to start from means no shape to guess an end from either, so
        # this is the whole document -- including the section before the projects.
        doc = spans(projects_dir, {"source": PROJEKTLISTE_NAME, "format": "docx", "title": "Alles"})
        texts = json.dumps(doc.sections[0].model_dump(mode="json"))
        assert BEFORE_HEADING in texts
        assert AFTER_SECTION in texts

    def test_omitting_the_begin_still_honours_an_end(self, projects_dir: Path) -> None:
        doc = spans(
            projects_dir,
            {
                "source": PROJEKTLISTE_NAME,
                "format": "docx",
                "end": "Projekthistorie",
                "title": "Vorab",
            },
        )
        texts = json.dumps(doc.sections[0].model_dump(mode="json"))
        assert BEFORE_HEADING in texts
        assert AFTER_SECTION not in texts

    def test_without_a_begin_no_heading_is_shown(self, projects_dir: Path) -> None:
        # There is no heading to name the section after, and the source
        # file's own filename is not a title -- it names the file, not what
        # is in it -- so no heading is shown for the section at all.
        doc = spans(
            projects_dir,
            {"source": PROJEKTLISTE_NAME, "format": "docx", "end": "Projekthistorie"},
        )
        assert doc.sections[-1].title is None

    def test_the_slug_still_falls_back_to_the_file_stem(self, projects_dir: Path) -> None:
        # An anchor still has to point *somewhere*, even with no heading shown.
        doc = spans(
            projects_dir,
            {"source": PROJEKTLISTE_NAME, "format": "docx", "end": "Projekthistorie"},
        )
        assert doc.sections[-1].slug == slugify(Path(PROJEKTLISTE_NAME).stem)

    def test_a_missing_begin_headline_is_an_error(self, projects_dir: Path) -> None:
        with pytest.raises(DocParseError, match="no heading matching 'Publikationen' found"):
            spans(
                projects_dir,
                {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Publikationen"},
            )

    def test_an_end_headline_that_never_comes_is_an_error(self, projects_dir: Path) -> None:
        with pytest.raises(DocParseError, match="no heading matching 'Publikationen' follows"):
            spans(
                projects_dir,
                {
                    "source": PROJEKTLISTE_NAME,
                    "format": "docx",
                    "begin": "Projekthistorie",
                    "end": "Publikationen",
                },
            )

    def test_a_missing_project_list_is_an_error_naming_the_entry(self, tmp_path: Path) -> None:
        # A recipe has many entries, so a message has to say which one failed.
        (tmp_path / "document.md").write_text(PROJECTS_MD, encoding="utf-8")
        config = {
            "sections": [
                {"source": "document.md", "format": "md", "end": "Projekte"},
                {"source": "*Projektliste*.docx", "format": "docx", "begin": "Projekthistorie"},
            ],
        }
        with pytest.raises(DocParseError, match=r"sections\[1\]: no file in .* matches"):
            build(tmp_path, config)

    def test_two_project_lists_are_an_error(self, tmp_path: Path) -> None:
        write_projektliste(tmp_path / "Projektliste_alt.docx")
        write_projektliste(tmp_path / "Projektliste_neu.docx")
        (tmp_path / "document.md").write_text(PROJECTS_MD, encoding="utf-8")
        config = {
            "sections": [
                {"source": "document.md", "format": "md", "end": "Projekte"},
                {"source": "*Projektliste*.docx", "format": "docx", "begin": "Projekthistorie"},
            ],
        }
        with pytest.raises(DocParseError, match="keep exactly one"):
            build(tmp_path, config)


# The rectangle every `TestXlsxSpans` entry names -- a bold, ruled header row
# and one data row, exercising a date and a currency figure. See
# `write_workbook`, which builds the sheet it comes from.
XLSX_RANGE: dict[str, Any] = {
    "col-start": "C",
    "col-end": "G",
    "row-start": 3,
    "row-end": 5,
}


class TestXlsxSpans:
    """An `.xlsx` entry brings in one cell rectangle as blocks, not Markdown."""

    @pytest.fixture
    def workbook_dir(self, tmp_path: Path) -> Path:
        write_workbook(tmp_path / "Rechnung.xlsx")
        (tmp_path / "document.md").write_text(PROJECTS_MD, encoding="utf-8")
        return tmp_path

    def test_the_section_carries_one_table(self, workbook_dir: Path) -> None:
        doc = spans(
            workbook_dir,
            {
                "source": "Rechnung.xlsx",
                "format": "xlsx",
                "title": "Rechnungsbeträge",
                **XLSX_RANGE,
            },
        )
        section = doc.section("rechnungsbetrage")
        assert section is not None
        assert [type(block) for block in section.blocks] == [RichTable]
        assert section.markdown == ""

    def test_the_source_file_is_recorded(self, workbook_dir: Path) -> None:
        doc = spans(
            workbook_dir,
            {
                "source": "Rechnung.xlsx",
                "format": "xlsx",
                "title": "Rechnungsbeträge",
                **XLSX_RANGE,
            },
        )
        section = doc.section("rechnungsbetrage")
        assert section is not None
        assert section.source is not None
        assert section.source.endswith("Rechnung.xlsx")

    def test_noframes_suppresses_the_source_s_own_borders(self, workbook_dir: Path) -> None:
        # XLSX_RANGE (rows 3-5) is the ruled range write_workbook produces --
        # see test_xlsx_import.py::TestFormatting::test_a_ruled_range_says_so.
        doc = spans(
            workbook_dir,
            {"source": "Rechnung.xlsx", "format": "xlsx", "noframes": True, **XLSX_RANGE},
        )
        section = doc.sections[-1]
        assert isinstance(section.blocks[0], RichTable)
        assert section.blocks[0].bordered is False

    def test_without_noframes_the_source_s_own_borders_still_apply(
        self, workbook_dir: Path
    ) -> None:
        doc = spans(
            workbook_dir,
            {"source": "Rechnung.xlsx", "format": "xlsx", **XLSX_RANGE},
        )
        section = doc.sections[-1]
        assert isinstance(section.blocks[0], RichTable)
        assert section.blocks[0].bordered is True

    def test_without_a_title_no_heading_is_shown(self, workbook_dir: Path) -> None:
        # There is no `begin` heading to fall back on, unlike a `.docx` entry:
        # row and column bounds carry no heading to name the section after
        # either, and the source file's own filename is not a title -- so no
        # heading is shown for the section at all.
        doc = spans(workbook_dir, {"source": "Rechnung.xlsx", "format": "xlsx", **XLSX_RANGE})
        assert doc.sections[-1].title is None

    def test_the_slug_still_falls_back_to_the_file_stem(self, workbook_dir: Path) -> None:
        doc = spans(workbook_dir, {"source": "Rechnung.xlsx", "format": "xlsx", **XLSX_RANGE})
        assert doc.sections[-1].slug == "rechnung"

    def test_a_missing_workbook_is_an_error_naming_the_entry(self, tmp_path: Path) -> None:
        (tmp_path / "document.md").write_text(PROJECTS_MD, encoding="utf-8")
        config = {
            "sections": [
                {"source": "document.md", "format": "md", "end": "Projekte"},
                {
                    "source": "Rechnung.xlsx",
                    "format": "xlsx",
                    "title": "Rechnungsbeträge",
                    **XLSX_RANGE,
                },
            ],
        }
        with pytest.raises(DocParseError, match=r"sections\[1\]: no such file"):
            build(tmp_path, config)


class TestTheHeader:
    """No `metadata` key: the header rides along with a span that starts at the top."""

    def test_frontmatter_and_summary_come_from_that_file(self, four_sections: Path) -> None:
        doc = spans(four_sections, {"source": "document.md", "format": "md", "begin": "Sprachen"})
        assert doc.name == "Ada Lovelace"
        assert doc.summary == "Ein Kurzprofil."

    def test_a_span_that_carries_sections_supplies_it_too(self, four_sections: Path) -> None:
        # The ordinary case: one entry brings the header *and* the first sections.
        doc = build(
            four_sections,
            {"sections": [{"source": "document.md", "format": "md", "end": "Kenntnisse"}]},
        )
        assert doc.name == "Ada Lovelace"
        assert doc.summary == "Ein Kurzprofil."
        assert [s.title for s in doc.sections] == ["Berufserfahrung", "Projekte"]

    def test_the_first_such_entry_wins(self, tmp_path: Path) -> None:
        # The Document has one name and one summary; the first entry to offer them is
        # the one that does, and a later file contributes only its sections.
        (tmp_path / "document.md").write_text(FOUR_SECTIONS, encoding="utf-8")
        (tmp_path / "other.md").write_text(
            "---\nname: Grace Hopper\n---\n\nEin anderes Profil.\n\n## Hobbys\n\n- Segeln\n",
            encoding="utf-8",
        )
        doc = build(
            tmp_path,
            {
                "sections": [
                    {"source": "document.md", "format": "md", "end": "Berufserfahrung"},
                    {"source": "other.md", "format": "md"},
                ]
            },
        )
        assert doc.name == "Ada Lovelace"
        assert doc.summary == "Ein Kurzprofil."
        assert [s.title for s in doc.sections] == ["Hobbys"]

    def test_a_docx_first_entry_does_not_supply_it(self, projects_dir: Path) -> None:
        # A Word document has no frontmatter, so the header comes from the later
        # Markdown entry instead of the recipe failing.
        doc = build(
            projects_dir,
            {
                "sections": [
                    {
                        "source": PROJEKTLISTE_NAME,
                        "format": "docx",
                        "begin": "Projekthistorie",
                        "title": "Projekte",
                    },
                    {"source": "document.md", "format": "md", "begin": "Kenntnisse"},
                    {"source": "document.md", "format": "md", "end": "Berufserfahrung"},
                ]
            },
        )
        assert doc.name == "Ada Lovelace"

    def test_no_entry_starting_at_the_top_gives_a_bare_header(self, four_sections: Path) -> None:
        # No Markdown entry has to supply a header at all: a recipe built purely
        # from spans that start mid-file (or, elsewhere, from `.docx`/`.xlsx`
        # only) still produces a document, just with no identity of its own.
        doc = build(
            four_sections,
            {"sections": [{"source": "document.md", "format": "md", "begin": "Sprachen"}]},
        )
        assert doc.name is None
        assert doc.summary is None
        assert [s.title for s in doc.sections] == ["Sprachen"]

    def test_the_default_output_name_is_used_with_no_header(self, four_sections: Path) -> None:
        loaded = load_doc(
            write_config(
                four_sections,
                {"sections": [{"source": "document.md", "format": "md", "begin": "Sprachen"}]},
            )
        )
        assert loaded.name == "document"

    def test_the_photo_is_relative_to_that_file(self, tmp_path: Path) -> None:
        sub = tmp_path / "sub"
        sub.mkdir()
        (sub / "portrait.png").write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 16)
        (sub / "document.md").write_text(
            "---\nname: Ada\nphoto: portrait.png\n---\n\n## Sprachen\n\n- Deutsch\n",
            encoding="utf-8",
        )
        doc = build(tmp_path, {"sections": [{"source": "sub/document.md", "format": "md"}]})
        assert doc.photo is not None

    def test_that_file_must_carry_frontmatter(self, tmp_path: Path) -> None:
        (tmp_path / "document.md").write_text("## Sprachen\n\n- Deutsch\n", encoding="utf-8")
        with pytest.raises(DocParseError, match="supplies the Document's header -- missing YAML"):
            build(tmp_path, {"sections": [{"source": "document.md", "format": "md"}]})


class TestNoMarkdownSource:
    """A recipe built entirely from `.docx`/`.xlsx` needs no Markdown at all."""

    def test_a_docx_and_xlsx_only_recipe_builds(self, projects_dir: Path) -> None:
        # No `document.md` is even referenced -- an invoice-shaped recipe.
        write_workbook(projects_dir / "Betraege.xlsx")
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"},
                    {"source": "Betraege.xlsx", "format": "xlsx", **XLSX_RANGE},
                ]
            },
        )
        assert doc.name is None
        assert doc.summary is None
        # The xlsx entry has no `begin` to fall back on and no `title` either,
        # so it shows no heading; the docx entry's `begin` still names one.
        assert [s.title for s in doc.sections] == ["Projekthistorie", None]
        assert [s.slug for s in doc.sections] == ["projekthistorie", "betraege"]

    def test_the_output_name_defaults_to_document(self, projects_dir: Path) -> None:
        loaded = load_doc(
            write_config(
                projects_dir,
                {
                    "sections": [
                        {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"}
                    ]
                },
            )
        )
        assert loaded.name == "document"

    def test_output_still_overrides_the_default_name(self, projects_dir: Path) -> None:
        loaded = load_doc(
            write_config(
                projects_dir,
                {
                    "sections": [
                        {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"}
                    ],
                    "output": "invoice",
                },
            )
        )
        assert loaded.name == "invoice"


class TestConfigPhoto:
    """The root-level `photo` key, an alternative to frontmatter's `photo:`."""

    def test_supplies_a_photo_with_no_markdown_source(
        self, projects_dir: Path, portrait_path: Path
    ) -> None:
        (projects_dir / "portrait.png").write_bytes(portrait_path.read_bytes())
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"}
                ],
                "photo": "portrait.png",
            },
        )
        assert doc.photo is not None
        assert doc.photo.media_type == "image/png"

    def test_overrides_a_photo_the_header_already_carries(
        self, tmp_path: Path, portrait_path: Path
    ) -> None:
        (tmp_path / "from-header.png").write_bytes(portrait_path.read_bytes())
        (tmp_path / "from-config.png").write_bytes(portrait_path.read_bytes())
        (tmp_path / "document.md").write_text(
            "---\nname: Ada\nphoto: from-header.png\n---\n\n## Sprachen\n\n- Deutsch\n",
            encoding="utf-8",
        )
        doc = build(
            tmp_path,
            {
                "sections": [{"source": "document.md", "format": "md"}],
                "photo": "from-config.png",
            },
        )
        assert doc.photo is not None
        assert doc.photo.data == (tmp_path / "from-config.png").read_bytes()

    def test_a_missing_photo_is_a_clean_error(self, projects_dir: Path) -> None:
        with pytest.raises(DocParseError, match="no such file"):
            build(
                projects_dir,
                {
                    "sections": [
                        {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"}
                    ],
                    "photo": "gone.png",
                },
            )


class TestSuppressedFilenameTitle:
    """A `.docx`/`.xlsx` section never shows a heading made of the source
    file's own filename -- see `_copy_docx`/`_copy_xlsx`. Position in the
    document does not matter: a filename is not a title in the middle of the
    document any more than it is at the top."""

    def test_a_docx_entry_with_no_title_or_begin_has_no_title(self, projects_dir: Path) -> None:
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": PROJEKTLISTE_NAME, "format": "docx", "end": "Projekthistorie"}
                ]
            },
        )
        assert doc.sections[0].title is None

    def test_the_slug_still_comes_from_the_file_stem(self, projects_dir: Path) -> None:
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": PROJEKTLISTE_NAME, "format": "docx", "end": "Projekthistorie"}
                ]
            },
        )
        assert doc.sections[0].slug == slugify(Path(PROJEKTLISTE_NAME).stem)

    def test_an_xlsx_entry_with_no_title_has_no_title(self, tmp_path: Path) -> None:
        write_workbook(tmp_path / "Rechnung.xlsx")
        doc = build(
            tmp_path, {"sections": [{"source": "Rechnung.xlsx", "format": "xlsx", **XLSX_RANGE}]}
        )
        assert doc.sections[0].title is None

    def test_a_later_docx_entry_with_no_title_or_begin_also_has_no_title(
        self, projects_dir: Path
    ) -> None:
        # Not just the document's first section -- a filename is never a
        # title, wherever the entry sits in the recipe.
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": "document.md", "format": "md"},
                    {"source": PROJEKTLISTE_NAME, "format": "docx", "end": "Projekthistorie"},
                ]
            },
        )
        assert doc.sections[-1].title is None

    def test_an_explicit_title_is_still_shown(self, projects_dir: Path) -> None:
        doc = build(
            projects_dir,
            {
                "sections": [
                    {
                        "source": PROJEKTLISTE_NAME,
                        "format": "docx",
                        "end": "Projekthistorie",
                        "title": "Projekte",
                    }
                ]
            },
        )
        assert doc.sections[0].title == "Projekte"

    def test_a_begin_derived_title_is_still_shown(self, projects_dir: Path) -> None:
        # `begin` names an actual heading in the source, not the filename --
        # only a bare filename is suppressed.
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": PROJEKTLISTE_NAME, "format": "docx", "begin": "Projekthistorie"}
                ]
            },
        )
        assert doc.sections[0].title == "Projekthistorie"

    def test_a_markdown_section_keeps_its_heading(self, four_sections: Path) -> None:
        # Only a filename-derived title is suppressed; a real `##` heading,
        # Markdown's own, is never touched by this.
        doc = build(four_sections, {"sections": [{"source": "document.md", "format": "md"}]})
        assert doc.sections[0].title == "Berufserfahrung"


class TestDocxHeaderFooter:
    """If the recipe's first entry is a `.docx`, its own page header and
    footer become the target document's page header and footer, exclusively
    -- see `cv_generator.docx_import.load_header`/`load_footer`."""

    @staticmethod
    def _write_docx_with_header_footer(
        path: Path, *, header_text: str = "Header text", footer_text: str = "Footer text"
    ) -> None:
        import docx

        document = docx.Document()
        document.add_paragraph("Body content.")
        section = document.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = header_text
        section.footer.is_linked_to_previous = False
        section.footer.paragraphs[0].text = footer_text
        document.save(str(path))

    def test_the_first_entry_s_header_and_footer_are_carried_over(self, projects_dir: Path) -> None:
        self._write_docx_with_header_footer(projects_dir / "Brief.docx")
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": "Brief.docx", "format": "docx"},
                    {"source": "document.md", "format": "md"},
                ]
            },
        )
        assert len(doc.page_header) == 1
        assert doc.page_header[0].text() == "Header text"  # type: ignore[union-attr]
        assert len(doc.page_footer) == 1
        assert doc.page_footer[0].text() == "Footer text"  # type: ignore[union-attr]

    def test_a_later_docx_entry_is_not_the_source(self, projects_dir: Path) -> None:
        # Only the recipe's first entry counts, not merely the first `.docx`
        # one -- and not whichever entry happens to carry a header or footer.
        self._write_docx_with_header_footer(projects_dir / "Brief.docx")
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": "document.md", "format": "md"},
                    {"source": "Brief.docx", "format": "docx"},
                ]
            },
        )
        assert doc.page_header == []
        assert doc.page_footer == []

    def test_a_docx_with_no_header_or_footer_content_leaves_them_empty(
        self, projects_dir: Path
    ) -> None:
        doc = build(
            projects_dir,
            {
                "sections": [
                    {
                        "source": PROJEKTLISTE_NAME,
                        "format": "docx",
                        "begin": "Projekthistorie",
                        "title": "Projekte",
                    },
                    {"source": "document.md", "format": "md", "begin": "Kenntnisse"},
                ]
            },
        )
        assert doc.page_header == []
        assert doc.page_footer == []

    def test_an_xlsx_first_entry_has_neither(self, projects_dir: Path) -> None:
        self._write_docx_with_header_footer(projects_dir / "Brief.docx")
        write_workbook(projects_dir / "Rechnung.xlsx")
        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": "Rechnung.xlsx", "format": "xlsx", **XLSX_RANGE},
                    {"source": "Brief.docx", "format": "docx"},
                ]
            },
        )
        assert doc.page_header == []
        assert doc.page_footer == []

    def test_a_background_image_behind_the_header_text_is_not_carried_over(
        self, projects_dir: Path
    ) -> None:
        # A letterhead-style header floats its logo as a full-page background
        # image behind the text (`wp:anchor behindDoc="1"`). Carrying it over
        # like a real picture would blow the header up to page size and push
        # the rest of the document onto a spurious extra page -- only the
        # text belongs in the imported header.
        import io

        import docx
        from docx.oxml.ns import qn

        document = docx.Document()
        document.add_paragraph("Body content.")
        header = document.sections[0].header
        header.is_linked_to_previous = False
        header.paragraphs[0].text = "Header text"
        picture_run = header.paragraphs[0].add_run()
        picture_run.add_picture(io.BytesIO(png(20, 20)))
        drawing = picture_run._r.find(qn("w:drawing"))
        inline = drawing.find(qn("wp:inline"))
        inline.tag = qn("wp:anchor")
        inline.set("behindDoc", "1")
        inline.set("simplePos", "0")
        path = projects_dir / "Brief.docx"
        document.save(str(path))

        doc = build(
            projects_dir,
            {
                "sections": [
                    {"source": "Brief.docx", "format": "docx"},
                    {"source": "document.md", "format": "md"},
                ]
            },
        )
        assert doc.page_header[0].text() == "Header text"  # type: ignore[union-attr]
        assert not any(run.image is not None for block in doc.page_header for run in block.runs)  # type: ignore[union-attr]


class TestUnsupportedSources:
    def test_an_unrecognised_format_value_is_rejected_by_the_schema(
        self, four_sections: Path
    ) -> None:
        # "format" is not guessed from the file, so a value neither reader
        # understands has to be caught here rather than surfacing as a read error.
        with pytest.raises(DocParseError, match="format"):
            load_config(
                write_config(
                    four_sections,
                    {"sections": [{"source": "document.md", "format": "txt", "begin": "Sprachen"}]},
                )
            )


class TestLoadDoc:
    def test_a_json_source_is_a_recipe(self, projects_path: Path) -> None:
        loaded = load_doc(projects_path)
        titles = [s.title for s in loaded.doc.sections]
        assert titles == ["Berufserfahrung", "Projekte", "Kenntnisse"]

    def test_a_markdown_source_is_a_whole_cv_on_its_own(self, minimal_path: Path) -> None:
        loaded = load_doc(minimal_path)
        assert loaded.doc.name == "Ada Lovelace"
        assert loaded.name == "minimal"

    def test_the_output_name_comes_from_the_recipe(self, projects_dir: Path) -> None:
        loaded = load_doc(write_config(projects_dir, PROJECTS_CONFIG | {"output": "lebenslauf"}))
        assert loaded.name == "lebenslauf"

    def test_it_otherwise_follows_the_file_the_header_came_from(self, projects_path: Path) -> None:
        # No `output` key in PROJECTS_CONFIG, and the header comes from document.md.
        assert load_doc(projects_path).name == "document"

    def test_the_recipe_s_own_name_is_not_the_output_name(self, projects_path: Path) -> None:
        # Otherwise every build would write dist/config.html.
        assert projects_path.stem == "config"
        assert load_doc(projects_path).name == "document"

    def test_sources_are_resolved_relative_to_the_recipe_not_the_cwd(
        self,
        projects_path: Path,
        tmp_path_factory: pytest.TempPathFactory,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        monkeypatch.chdir(tmp_path_factory.mktemp("elsewhere"))
        projects = load_doc(projects_path).doc.section("projekte")
        assert projects is not None
        assert projects.blocks

    def test_with_no_target_the_loaded_cv_names_none(self, projects_path: Path) -> None:
        assert load_doc(projects_path).target is None

    def test_target_is_resolved_relative_to_the_project_root(
        self, projects_dir: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # The recipe's own directory must not matter here -- `target` is a
        # project-root path, the same convention `source`'s fallback uses.
        monkeypatch.chdir(projects_dir)
        config_path = write_config(projects_dir, PROJECTS_CONFIG | {"target": "exports/doc"})
        loaded = load_doc(config_path)
        assert loaded.target == projects_dir / "exports" / "doc"

    def test_target_does_not_depend_on_the_recipe_s_directory(
        self,
        projects_dir: Path,
        tmp_path_factory: pytest.TempPathFactory,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        root = tmp_path_factory.mktemp("project-root")
        monkeypatch.chdir(root)
        config_path = write_config(projects_dir, PROJECTS_CONFIG | {"target": "exports/doc"})
        loaded = load_doc(config_path)
        assert loaded.target == root / "exports" / "doc"


class TestRepositoryRecipe:
    """`data/config.json` is the project's own example, so it has to work."""

    def test_it_assembles_the_sample_document(self, sample_config_path: Path) -> None:
        doc, name, target = load_doc(sample_config_path)
        assert name == "document"
        assert target is None
        assert [s.title for s in doc.sections] == [
            "Berufserfahrung",
            "Rechnungsbeträge",
            "Kenntnisse",
            "Ausbildung",
            "Sprachen",
        ]

    def test_the_invoice_figures_come_from_the_workbook(self, sample_config_path: Path) -> None:
        section = load_doc(sample_config_path).doc.section("rechnungsbetrage")
        assert section is not None
        assert section.blocks
        assert section.source is not None
        assert section.source.endswith(".xlsx")

    def test_every_other_section_comes_from_the_markdown_file(
        self, sample_config_path: Path
    ) -> None:
        doc = load_doc(sample_config_path).doc
        markdown = [s for s in doc.sections if s.slug != "rechnungsbetrage"]
        assert all(s.source is not None and s.source.endswith("document.md") for s in markdown)
        assert all(s.markdown for s in markdown)


def test_build_cv_takes_a_config_object(projects_dir: Path) -> None:
    # The API a caller assembling a recipe in Python uses; the CLI is one caller.
    config = BuildConfig.model_validate(PROJECTS_CONFIG)
    doc, name, target = build_doc(config, projects_dir)
    assert [s.title for s in doc.sections] == ["Berufserfahrung", "Projekte", "Kenntnisse"]
    assert name == "document"
    assert target is None


def test_build_cv_falls_back_to_the_project_root(tmp_path: Path) -> None:
    # The recipe lives in a subdirectory of the project; its sources live at the
    # project root instead of beside it -- resolve_source's project_root fallback
    # is what still finds them.
    write_projektliste(tmp_path / PROJEKTLISTE_NAME)
    (tmp_path / "document.md").write_text(PROJECTS_MD, encoding="utf-8")
    recipe_dir = tmp_path / "recipe"
    recipe_dir.mkdir()
    config = BuildConfig.model_validate(PROJECTS_CONFIG)

    doc, name, target = build_doc(config, recipe_dir, project_root=tmp_path)

    assert [s.title for s in doc.sections] == ["Berufserfahrung", "Projekte", "Kenntnisse"]
    assert name == "document"
    assert target is None
