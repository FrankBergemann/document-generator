"""Tests for `.docx` output.

Assertions go through python-docx reading the file back, plus raw XML where
there is no read API (hyperlinks, borders). A `.docx` that "looks fine" but has
mis-ordered XML gets silently repaired by Word, so the structural checks matter
as much as the visible content.
"""

from __future__ import annotations

import struct
import zlib
from itertools import pairwise
from pathlib import Path

import docx
import pytest
from docx.document import Document as WordDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.table import Table

from cv_generator.errors import RenderError
from cv_generator.models import (
    CV,
    Contact,
    Link,
    Photo,
    RichCell,
    RichParagraph,
    RichRun,
    RichTable,
    Section,
)
from cv_generator.ooxml import HYPERLINK_RELATIONSHIP
from cv_generator.parser import parse_cv
from cv_generator.word import (
    STYLE_BODY,
    STYLE_CODE,
    STYLE_CONTACT,
    STYLE_ENTRY,
    STYLE_HEADLINE,
    STYLE_META,
    STYLE_NAME,
    STYLE_QUOTE,
    STYLE_SECTION,
    STYLE_SUMMARY,
    WordRenderer,
    WordTheme,
)
from tests.support import BODY_SIZE_PT, SAMPLE_PROJECTS


def write(cv: CV, tmp_path: Path, name: str = "document.docx") -> Path:
    output = tmp_path / name
    WordRenderer().render(cv, output)
    return output


def read(path: Path) -> WordDocument:
    return docx.Document(str(path))


def png(width: int, height: int) -> bytes:
    """A minimal valid PNG of exactly this pixel size.

    Built by hand because the sizing rules need images of a chosen aspect ratio
    and the project has no image library to draw one with.
    """

    def chunk(kind: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + kind
            + payload
            + struct.pack(">I", zlib.crc32(kind + payload))
        )

    header = struct.pack(">2I5B", width, height, 8, 0, 0, 0, 0)  # 8-bit greyscale
    rows = b"".join(b"\x00" + b"\xff" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", header)
        + chunk(b"IDAT", zlib.compress(rows))
        + chunk(b"IEND", b"")
    )


def texts(document: WordDocument, style: str) -> list[str]:
    return [p.text for p in document.paragraphs if p.style is not None and p.style.name == style]


def styles_in_order(document: WordDocument) -> list[str]:
    return [p.style.name for p in document.paragraphs if p.style is not None]


@pytest.fixture
def minimal_doc(minimal_cv: CV, tmp_path: Path) -> WordDocument:
    return read(write(minimal_cv, tmp_path))


@pytest.fixture
def rich_doc(rich_cv: CV, tmp_path: Path) -> WordDocument:
    return read(write(rich_cv, tmp_path))


@pytest.fixture
def photo_doc(photo_cv: CV, tmp_path: Path) -> WordDocument:
    return read(write(photo_cv, tmp_path, "photo.docx"))


def header_table(document: WordDocument) -> Table:
    return document.tables[0]


class TestFileOutput:
    def test_writes_a_readable_docx(self, minimal_cv: CV, tmp_path: Path) -> None:
        output = write(minimal_cv, tmp_path)
        assert output.is_file()
        # A .docx is a zip; a corrupt one fails here rather than in Word.
        assert output.read_bytes().startswith(b"PK")
        assert read(output).paragraphs

    def test_creates_missing_directories(self, minimal_cv: CV, tmp_path: Path) -> None:
        output = tmp_path / "nested" / "deeper" / "document.docx"
        WordRenderer().render(minimal_cv, output)
        assert output.is_file()

    def test_unwritable_target_raises_cv_error(self, minimal_cv: CV, tmp_path: Path) -> None:
        # A directory where a file should go: the OSError must surface as a
        # CVError so the CLI reports it on one line instead of a traceback.
        blocked = tmp_path / "document.docx"
        blocked.mkdir()
        with pytest.raises(RenderError, match="cannot write"):
            WordRenderer().render(minimal_cv, blocked)


class TestPageSetup:
    # Word stores lengths in twips, so a value set in mm does not round-trip
    # exactly; a hundredth of a millimetre is well inside the noise.
    TOLERANCE_MM = 0.02

    def test_a4_portrait(self, minimal_doc: WordDocument) -> None:
        page = minimal_doc.sections[0]
        assert page.page_width is not None and page.page_height is not None
        assert page.page_width.mm == pytest.approx(210, abs=self.TOLERANCE_MM)
        assert page.page_height.mm == pytest.approx(297, abs=self.TOLERANCE_MM)

    def test_margins_match_the_theme(self, minimal_doc: WordDocument) -> None:
        page = minimal_doc.sections[0]
        theme = WordTheme()
        assert page.top_margin is not None and page.left_margin is not None
        assert page.top_margin.mm == pytest.approx(
            theme.page_margin_vertical_mm, abs=self.TOLERANCE_MM
        )
        assert page.left_margin.mm == pytest.approx(
            theme.page_margin_horizontal_mm, abs=self.TOLERANCE_MM
        )

    def test_metadata_names_the_person(self, minimal_doc: WordDocument) -> None:
        assert minimal_doc.core_properties.title == "Ada Lovelace"
        assert minimal_doc.core_properties.author == "Ada Lovelace"

    def test_language_is_tagged_for_spellcheck(self, rich_doc: WordDocument) -> None:
        properties = rich_doc.styles["Normal"].element.rPr
        assert properties is not None
        assert properties.find(qn("w:lang")).get(qn("w:val")) == "en"


class TestHeader:
    def test_name_and_headline(self, minimal_doc: WordDocument) -> None:
        assert texts(minimal_doc, STYLE_NAME) == ["Ada Lovelace"]
        assert texts(minimal_doc, STYLE_HEADLINE) == ["Mathematician"]

    def test_headline_omitted_when_absent(self, tmp_path: Path) -> None:
        document = read(write(CV(name="Ada"), tmp_path))
        assert texts(document, STYLE_HEADLINE) == []

    def test_contact_details_are_present(self, rich_doc: WordDocument) -> None:
        xml = rich_doc.element.xml
        for fragment in ("ada@example.com", "+44 1815 121815", "London", "Notes"):
            assert fragment in xml

    def test_email_and_links_become_hyperlinks(self, rich_cv: CV, tmp_path: Path) -> None:
        document = read(write(rich_cv, tmp_path))
        targets = {
            rel.target_ref
            for rel in document.part.rels.values()
            if rel.reltype == HYPERLINK_RELATIONSHIP
        }
        assert "mailto:ada@example.com" in targets
        assert "https://example.com/notes" in targets
        assert qn("w:hyperlink") in document.element.xml or "w:hyperlink" in document.element.xml

    def test_header_carries_a_rule(self, minimal_doc: WordDocument) -> None:
        # The rule sits on the last header paragraph, mirroring
        # `.cv-header { border-bottom: ... }`.
        contact = next(
            p
            for p in minimal_doc.paragraphs
            if p.style is not None and p.style.name == STYLE_CONTACT
        )
        properties = contact._p.find(qn("w:pPr"))
        assert properties is not None
        assert properties.find(qn("w:pBdr")) is not None

    def test_border_follows_the_paragraph_style_in_xml(self, minimal_doc: WordDocument) -> None:
        # w:pBdr must come after w:pStyle or Word repairs the file on open.
        properties = next(
            p._p.find(qn("w:pPr"))
            for p in minimal_doc.paragraphs
            if p.style is not None and p.style.name == STYLE_SECTION
        )
        tags = [child.tag for child in properties]
        assert tags.index(qn("w:pStyle")) < tags.index(qn("w:pBdr"))


class TestHeaderPhoto:
    """The photo turns the header into two columns; without one nothing changes."""

    def test_no_photo_means_no_layout_table(self, minimal_doc: WordDocument) -> None:
        assert minimal_doc.tables == []

    def test_header_becomes_a_single_row_two_column_table(self, photo_doc: WordDocument) -> None:
        # Word has no flexbox; a borderless layout table is how a column pair is
        # expressed. Its cells must not inherit the visible grid style.
        table = header_table(photo_doc)
        assert (len(table.rows), len(table.columns)) == (1, 2)
        # No cell borders: the only rule in this table is the header's own.
        assert "w:tcBorders" not in table._tbl.xml

    def test_identity_moves_into_the_left_cell(self, photo_doc: WordDocument) -> None:
        cell = header_table(photo_doc).rows[0].cells[0]
        assert [p.style.name for p in cell.paragraphs if p.style is not None] == [
            STYLE_NAME,
            STYLE_HEADLINE,
            STYLE_CONTACT,
        ]
        assert cell.paragraphs[0].text == "Ada Lovelace"

    def test_no_blank_line_above_the_name(self, photo_doc: WordDocument) -> None:
        # A fresh cell arrives with an empty paragraph; leaving it in place would
        # push the name down a line.
        assert header_table(photo_doc).rows[0].cells[0].paragraphs[0].text != ""

    def test_columns_fill_the_text_width(self, photo_doc: WordDocument) -> None:
        theme = WordTheme()
        widths = [column.width for column in header_table(photo_doc).columns]
        assert all(width is not None for width in widths)
        total = sum(width.mm for width in widths if width is not None)
        assert total == pytest.approx(theme.content_width_mm, abs=0.05)

    def test_photo_column_is_wide_enough_for_the_photo_and_its_gap(
        self, photo_doc: WordDocument
    ) -> None:
        theme = WordTheme()
        width = header_table(photo_doc).columns[1].width
        assert width is not None
        assert width.mm == pytest.approx(theme.photo_width_mm + theme.photo_gap_mm, abs=0.05)

    def test_photo_is_pushed_to_the_right_margin(self, photo_doc: WordDocument) -> None:
        cell = header_table(photo_doc).rows[0].cells[1]
        assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def test_cell_insets_keep_the_name_aligned_with_the_body(self, photo_doc: WordDocument) -> None:
        # Word's default left inset would indent the name relative to every
        # section below it.
        margins = header_table(photo_doc)._tbl.tblPr.find(qn("w:tblCellMar"))
        assert margins is not None
        assert margins.find(qn("w:left")).get(qn("w:w")) == "0"

    def test_the_rule_spans_both_columns(self, photo_doc: WordDocument) -> None:
        # A paragraph border would underline the contact line only, so the rule
        # mirroring `.cv-header { border-bottom }` lives on the table.
        borders = header_table(photo_doc)._tbl.tblPr.find(qn("w:tblBorders"))
        assert borders is not None
        assert borders.find(qn("w:bottom")).get(qn("w:color")) == WordTheme().accent

    def test_cell_inset_edges_stay_in_schema_order(self, photo_doc: WordDocument) -> None:
        # CT_TblCellMar declares top, left, bottom, right. Emitting them in any
        # other order is the kind of thing Word repairs silently on open.
        margins = header_table(photo_doc)._tbl.tblPr.find(qn("w:tblCellMar"))
        assert margins is not None
        assert [child.tag for child in margins] == [
            qn(f"w:{edge}") for edge in ("top", "left", "bottom", "right")
        ]

    def test_table_properties_stay_in_schema_order(self, photo_doc: WordDocument) -> None:
        # w:tblBorders before w:tblLayout before w:tblCellMar, or Word "repairs"
        # the document on open.
        tags = [child.tag for child in header_table(photo_doc)._tbl.tblPr]
        for earlier, later in (
            ("w:tblW", "w:tblBorders"),
            ("w:tblBorders", "w:tblLayout"),
            ("w:tblLayout", "w:tblCellMar"),
        ):
            assert tags.index(qn(earlier)) < tags.index(qn(later))

    def test_image_bytes_are_embedded(self, photo_doc: WordDocument, portrait_path: Path) -> None:
        blobs = {
            rel.target_part.blob for rel in photo_doc.part.rels.values() if rel.reltype == RT.IMAGE
        }
        assert portrait_path.read_bytes() in blobs

    def test_photo_is_as_wide_as_the_theme_says(self, photo_doc: WordDocument) -> None:
        shape = photo_doc.inline_shapes[0]
        assert shape.width.mm == pytest.approx(WordTheme().photo_width_mm, abs=0.05)

    def test_aspect_ratio_is_preserved(self, photo_doc: WordDocument) -> None:
        # The fixture is 120x160 px, so 34mm wide must come out 45.3mm tall.
        shape = photo_doc.inline_shapes[0]
        assert shape.height / shape.width == pytest.approx(160 / 120, rel=1e-3)

    def test_a_tall_photo_is_bounded_without_being_squashed(self, tmp_path: Path) -> None:
        # Mirrors `max-height` in the stylesheet: a 1:10 image scaled to 34mm
        # wide would be 340mm tall and push the whole CV off page one.
        theme = WordTheme()
        cv = CV(name="Ada", photo=Photo(data=png(100, 1000), media_type="image/png"))
        shape = read(write(cv, tmp_path, "tall.docx")).inline_shapes[0]
        assert shape.height.mm == pytest.approx(theme.photo_max_height_mm, abs=0.05)
        assert shape.height / shape.width == pytest.approx(10.0, rel=1e-3)

    def test_photo_survives_a_second_render(self, photo_cv: CV, tmp_path: Path) -> None:
        # The image stream is consumed on use; a reused renderer must not hand
        # an exhausted one to the second document.
        renderer = WordRenderer()
        renderer.render(photo_cv, tmp_path / "one.docx")
        renderer.render(photo_cv, tmp_path / "two.docx")
        assert len(read(tmp_path / "two.docx").inline_shapes) == 1

    def test_theme_can_resize_the_photo(self, photo_cv: CV, tmp_path: Path) -> None:
        output = tmp_path / "narrow.docx"
        WordRenderer(WordTheme(photo_width_mm=25.0)).render(photo_cv, output)
        assert read(output).inline_shapes[0].width.mm == pytest.approx(25.0, abs=0.05)

    def test_widening_the_photo_still_respects_the_height_bound(
        self, photo_cv: CV, tmp_path: Path
    ) -> None:
        # 50mm wide would make the 3:4 fixture 66.7mm tall, so the bound wins and
        # takes the width back down with it rather than distorting the picture.
        output = tmp_path / "wide.docx"
        WordRenderer(WordTheme(photo_width_mm=50.0)).render(photo_cv, output)
        shape = read(output).inline_shapes[0]
        assert shape.height.mm == pytest.approx(WordTheme().photo_max_height_mm, abs=0.05)
        assert shape.width.mm == pytest.approx(36.0, abs=0.05)


class TestSections:
    def test_section_titles_use_the_heading_style(self, minimal_doc: WordDocument) -> None:
        assert texts(minimal_doc, STYLE_SECTION) == ["Experience", "Skills"]

    def test_section_order_is_preserved(self, rich_doc: WordDocument) -> None:
        assert texts(rich_doc, STYLE_SECTION) == ["Experience", "Notes"]

    def test_summary_precedes_the_first_section(self, minimal_doc: WordDocument) -> None:
        order = styles_in_order(minimal_doc)
        assert order.index(STYLE_SUMMARY) < order.index(STYLE_SECTION)

    def test_summary_omitted_when_absent(self, tmp_path: Path) -> None:
        cv = parse_cv("---\nname: Ada\n---\n\n## Skills\n\n- Python\n")
        assert texts(read(write(cv, tmp_path)), STYLE_SUMMARY) == []


class TestEntries:
    def test_h3_becomes_an_entry(self, minimal_doc: WordDocument) -> None:
        assert texts(minimal_doc, STYLE_ENTRY) == ["Analyst — Analytical Engine"]

    def test_italic_line_under_an_entry_is_treated_as_meta(self, minimal_doc: WordDocument) -> None:
        assert texts(minimal_doc, STYLE_META) == ["1842 – 1843"]

    def test_a_normal_paragraph_is_not_meta(self, tmp_path: Path) -> None:
        cv = parse_cv("---\nname: Ada\n---\n## S\n\n*Just italic text.*\n")
        document = read(write(cv, tmp_path))
        assert texts(document, STYLE_META) == []
        assert "Just italic text." in texts(document, STYLE_BODY)

    def test_entry_heading_stays_with_its_content(self, minimal_doc: WordDocument) -> None:
        entry = next(p for p in minimal_doc.paragraphs if p.text.startswith("Analyst"))
        assert entry.style is not None
        assert entry.style.paragraph_format.keep_with_next is True


class TestBlockContent:
    def test_bullets_use_word_list_styles(self, minimal_doc: WordDocument) -> None:
        assert "Wrote Note G." in texts(minimal_doc, "List Bullet")

    def test_nested_bullets_indent_one_level_deeper(self, rich_doc: WordDocument) -> None:
        assert texts(rich_doc, "List Bullet 2") == ["Sub-point one.", "Sub-point two."]

    def test_ordered_lists_are_numbered(self, rich_doc: WordDocument) -> None:
        assert texts(rich_doc, "List Number") == ["First ordered step.", "Second ordered step."]

    def test_blockquote(self, rich_doc: WordDocument) -> None:
        assert texts(rich_doc, STYLE_QUOTE) == ["Quoted remark about the engine."]

    def test_fenced_code_keeps_its_lines(self, rich_doc: WordDocument) -> None:
        code = texts(rich_doc, STYLE_CODE)
        assert len(code) == 1
        assert "def bernoulli(n):" in code[0]
        assert "return n" in code[0]

    def test_tables_become_word_tables(self, rich_doc: WordDocument) -> None:
        assert len(rich_doc.tables) == 1
        table = rich_doc.tables[0]
        assert len(table.rows) == 3
        assert [cell.text for cell in table.rows[0].cells] == ["Machine", "Year"]
        assert table.rows[1].cells[0].text == "Difference Engine"

    def test_table_header_is_bold(self, rich_doc: WordDocument) -> None:
        header = rich_doc.tables[0].rows[0].cells[0].paragraphs[0]
        assert header.runs[0].font.bold is True

    def test_horizontal_rule_and_trailing_text_survive(self, rich_doc: WordDocument) -> None:
        assert "Closing paragraph after a rule." in texts(rich_doc, STYLE_BODY)

    def test_no_markdown_syntax_leaks_into_the_document(self, rich_doc: WordDocument) -> None:
        xml = rich_doc.element.xml
        for leak in ("**bold**", "~~struck~~", "<li>", "<strong>"):
            assert leak not in xml


class TestInlineFormatting:
    @pytest.fixture
    def summary_runs(self, rich_doc: WordDocument) -> dict[str, object]:
        paragraph = next(
            p for p in rich_doc.paragraphs if p.style is not None and p.style.name == STYLE_SUMMARY
        )
        return {run.text: run for run in paragraph.runs}

    def test_bold(self, summary_runs: dict[str, object]) -> None:
        assert summary_runs["bold"].font.bold is True  # type: ignore[attr-defined]

    def test_italic(self, summary_runs: dict[str, object]) -> None:
        assert summary_runs["italic"].font.italic is True  # type: ignore[attr-defined]

    def test_strikethrough(self, summary_runs: dict[str, object]) -> None:
        assert summary_runs["struck"].font.strike is True  # type: ignore[attr-defined]

    def test_inline_code_uses_the_mono_font(self, summary_runs: dict[str, object]) -> None:
        assert summary_runs["code"].font.name == WordTheme().mono_font  # type: ignore[attr-defined]

    def test_inline_link_is_a_real_hyperlink(self, rich_cv: CV, tmp_path: Path) -> None:
        document = read(write(rich_cv, tmp_path))
        targets = {
            rel.target_ref
            for rel in document.part.rels.values()
            if rel.reltype == HYPERLINK_RELATIONSHIP
        }
        assert "https://example.com/summary" in targets

    def test_softbreak_becomes_a_space_not_a_new_paragraph(self, rich_doc: WordDocument) -> None:
        summary = texts(rich_doc, STYLE_SUMMARY)
        assert len(summary) == 1
        assert "and a" in summary[0]

    def test_run_formatting_does_not_unset_style_bold(self, minimal_doc: WordDocument) -> None:
        # Runs only ever switch attributes on; setting bold=False explicitly
        # would override the entry style and un-bold every heading.
        entry = next(p for p in minimal_doc.paragraphs if p.text.startswith("Analyst"))
        assert all(run.font.bold is not False for run in entry.runs)


class TestImportedBlocks:
    """A section imported from a .docx keeps that document's formatting."""

    @pytest.fixture
    def projects_doc(self, projects_cv: CV, tmp_path: Path) -> WordDocument:
        return read(write(projects_cv, tmp_path, "projects.docx"))

    def test_each_project_is_a_table(self, projects_doc: WordDocument) -> None:
        assert len(projects_doc.tables) == len(SAMPLE_PROJECTS)

    def test_adjacent_tables_are_separated_by_a_paragraph(self, projects_doc: WordDocument) -> None:
        # Word merges two `w:tbl` siblings with nothing between them into one
        # table -- the blank paragraph the source used to keep them apart is
        # dropped on import, so it has to be reinstated here, not carried over.
        body = projects_doc.element.body
        tables = [child for child in body if child.tag == qn("w:tbl")]
        assert len(tables) >= 2
        for earlier, later in pairwise(tables):
            between = list(earlier.itersiblings(preceding=False))
            assert later in between
            assert any(sibling.tag == qn("w:p") for sibling in between[: between.index(later)])

    def test_cell_content_survives_paragraph_by_paragraph(self, projects_doc: WordDocument) -> None:
        period = projects_doc.tables[0].rows[0].cells[0]
        assert [p.text for p in period.paragraphs] == [
            "02/2026 – 07/2026",
            "Rolle:",
            "programmweiter Testmanager",
            "",
            "Kunde:",
            "Land Schleswig-Holstein",
        ]

    def test_bold_and_size_come_from_the_source_document(self, projects_doc: WordDocument) -> None:
        period = projects_doc.tables[0].rows[0].cells[0].paragraphs[0]
        assert period.runs[0].font.bold is True
        assert period.runs[0].font.size is not None
        assert period.runs[0].font.size.pt == BODY_SIZE_PT

    def test_bullets_use_word_list_styles(self, projects_doc: WordDocument) -> None:
        detail = projects_doc.tables[0].rows[0].cells[1]
        assert [p.style.name for p in detail.paragraphs if p.style is not None] == [
            STYLE_BODY,
            "List Bullet",
            "List Bullet",
            "List Bullet 2",
            STYLE_BODY,
        ]

    def test_a_ruled_source_table_is_ruled_here(self, projects_doc: WordDocument) -> None:
        table = projects_doc.tables[0]
        assert table.style is not None
        assert table.style.name == "Table Grid"

    def test_column_widths_are_scaled_to_this_page(self, projects_doc: WordDocument) -> None:
        widths = [column.width for column in projects_doc.tables[0].columns]
        assert all(width is not None for width in widths)
        content_mm = WordTheme().content_width_mm
        assert [round(width.mm) for width in widths if width is not None] == [
            round(content_mm * 0.2),
            round(content_mm * 0.8),
        ]

    def test_hyperlink_is_a_real_hyperlink(self, projects_doc: WordDocument) -> None:
        targets = [
            rel.target_ref
            for rel in projects_doc.part.rels.values()
            if rel.reltype == HYPERLINK_RELATIONSHIP
        ]
        assert "https://example.org/projekt" in targets

    def test_the_section_heading_stays_this_cv_s_own(self, projects_doc: WordDocument) -> None:
        # Only the section's *content* is imported; the heading is the CV's.
        assert "Projekte" in texts(projects_doc, STYLE_SECTION)

    def test_markdown_body_of_the_section_is_absent(self, projects_doc: WordDocument) -> None:
        assert "darf nicht" not in projects_doc.element.xml

    def test_every_cell_keeps_a_paragraph(self, projects_doc: WordDocument) -> None:
        # Word repairs a table cell that contains no paragraph at all.
        for table in projects_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    assert cell.paragraphs


def _cv_with_table(*, centered: bool, column_widths: list[float] | None = None) -> CV:
    cell = RichCell(paragraphs=[RichParagraph(runs=[RichRun(text="x")])])
    table = RichTable(rows=[[cell, cell]], centered=centered, column_widths=column_widths or [])
    section = Section(title="Zahlen", slug="zahlen", markdown="", blocks=[table])
    return CV(name="Ada", sections=[section])


class TestCenteredImportedTable:
    """An `.xlsx` import sizes its table to its content and centers it."""

    def test_a_centered_table_is_centered_and_autofits(self, tmp_path: Path) -> None:
        doc = read(write(_cv_with_table(centered=True), tmp_path))
        table = doc.tables[0]
        assert table.alignment == WD_TABLE_ALIGNMENT.CENTER
        assert table.autofit is True

    def test_a_non_centered_table_is_scaled_to_the_page_instead(self, tmp_path: Path) -> None:
        doc = read(write(_cv_with_table(centered=False, column_widths=[0.5, 0.5]), tmp_path))
        table = doc.tables[0]
        assert table.alignment is None
        widths = [column.width for column in table.columns]
        assert all(width is not None for width in widths)


class TestTheme:
    def test_defaults_mirror_the_html_theme(self) -> None:
        theme = WordTheme()
        assert theme.accent == "2F5D8A"
        assert theme.body_size == 10.5

    def test_theme_is_applied_to_styles(self, minimal_cv: CV, tmp_path: Path) -> None:
        renderer = WordRenderer(WordTheme(body_font="Georgia", body_size=12.0))
        output = tmp_path / "themed.docx"
        renderer.render(minimal_cv, output)
        normal = read(output).styles["Normal"]
        assert normal.font.name == "Georgia"
        assert normal.font.size is not None
        assert normal.font.size.pt == 12.0

    def test_renderer_is_reusable(self, minimal_cv: CV, tmp_path: Path) -> None:
        # Styles are added per document; a shared renderer must not trip over
        # "style already exists" on its second call.
        renderer = WordRenderer()
        renderer.render(minimal_cv, tmp_path / "one.docx")
        renderer.render(minimal_cv, tmp_path / "two.docx")
        assert (tmp_path / "two.docx").is_file()


class TestEscaping:
    def test_xml_special_characters_are_handled(self, tmp_path: Path) -> None:
        cv = CV(
            name="Ada & <Lovelace>",
            contact=Contact(links=[Link(label="A & B", url="https://x.test/?a=1&b=2")]),
        )
        document = read(write(cv, tmp_path))
        assert texts(document, STYLE_NAME) == ["Ada & <Lovelace>"]
