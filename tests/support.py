"""Shared test helpers.

Chromium is a ~150 MB optional download, so the suite has to pass without it --
but it must also really exercise the engine when it is present, rather than
skipping silently forever. "Present" means either installed locally or served by
a browser container (``CV_GENERATOR_BROWSER_WS``); ``requires_chromium`` covers
both, so the same tests run against a remote browser in the dev container.

Probing happens once per session here: the local probe costs a Playwright driver
start-up and the remote one a TCP connect.

The project-list fixture is *built* rather than committed as a binary: a
``.docx`` in ``tests/data`` would be an opaque blob nobody can review or amend,
while :func:`write_projektliste` says in Python exactly which formatting the
importer is expected to read back.
"""

from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass, field
from pathlib import Path

import docx
import pytest
from docx.document import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from cv_generator.pdf.chrome import ChromeEngine, local_browser_installed

DATA_DIR = Path(__file__).parent / "data"
REPO_ROOT = Path(__file__).parent.parent

CHROME_AVAILABLE = ChromeEngine().is_available()

# Tracked separately because a test about the *missing browser* message must not
# run on a machine that has a browser but a stopped browser container.
LOCAL_CHROME_INSTALLED = local_browser_installed()

requires_chromium = pytest.mark.skipif(
    not CHROME_AVAILABLE,
    reason=(
        "no chromium available (pip install '.[pdf]' && playwright install chromium, "
        "or start the playwright service)"
    ),
)

# -- the Word project list ------------------------------------------------

PROJEKTLISTE_NAME = "Bergemann-Projektliste_01_01_2026.docx"

# Matches the real document: headings are bold 12pt with no heading style at all,
# body text is 10pt, and each project is a two-column table.
HEADING_SIZE_PT = 12.0
BODY_SIZE_PT = 10.0
PERIOD_COLUMN_MM = 30.0
DETAIL_COLUMN_MM = 120.0

BEFORE_HEADING = "Vor der Projekthistorie"
AFTER_SECTION = "Promotion Maschinenbau"

BOLD_STYLE = "Fixture Bold"
INHERITS_BOLD = "bold from the style"
SWITCHED_OFF = " and off again"
SPLIT_RUNS = ("Ein Satz, ", "in zwei Runs")


@dataclass(frozen=True)
class Project:
    """One project, as the project list lays it out."""

    period: str
    role: str
    customer: str
    activities: Sequence[str] = field(default_factory=tuple)
    # Rendered one level deeper, to cover nested bullets.
    details: Sequence[str] = field(default_factory=tuple)
    link: tuple[str, str] | None = None


SAMPLE_PROJECTS: tuple[Project, ...] = (
    Project(
        period="02/2026 – 07/2026",
        role="programmweiter Testmanager",
        customer="Land Schleswig-Holstein",
        activities=("Konzeption einer Integrationstestumgebung", "Aufbau der Testautomatisierung"),
        details=("mit Playwright & TypeScript",),
        link=("Projektseite", "https://example.org/projekt"),
    ),
    Project(
        period="05/2025 – 09/2025",
        role="SAP-Testmanager",
        customer="Louis Motorrad GmbH",
        activities=("Testmanagement für S/4HANA",),
    ),
)


def write_projektliste(
    path: Path,
    projects: Sequence[Project] = SAMPLE_PROJECTS,
    *,
    heading: str = "Projekthistorie",
) -> Path:
    """Write a Word project list of the shape this project imports from.

    Around the section of interest it puts content that must *not* be imported:
    another section before it and one after it, both introduced by a heading
    formatted the same way, because that similarity is how the end is found.
    """
    document = docx.Document()
    _heading(document.add_paragraph(), "Profil")
    _bullet(document.add_paragraph(), BEFORE_HEADING)

    _heading(document.add_paragraph(), heading)
    document.add_paragraph()
    for project in projects:
        _project_table(document, project)
        # The blank paragraph Word needs between two tables, or they merge.
        document.add_paragraph()

    _heading(document.add_paragraph(), "Ausbildung")
    _text(document.add_paragraph(), AFTER_SECTION)

    document.save(str(path))
    return path


def write_styled_docx(path: Path, *, heading: str = "Projekthistorie") -> Path:
    """A document whose formatting only makes sense once styles are resolved.

    The first paragraph is in a bold style and holds a run that inherits the
    bold and a run that switches it off the way Word does (``<w:b w:val="0"/>``).
    Reading only the direct formatting gets both runs wrong, in opposite
    directions. The second is one sentence split across two runs of identical
    formatting, as Word leaves it after an edit.
    """
    document = docx.Document()
    style = document.styles.add_style(BOLD_STYLE, WD_STYLE_TYPE.PARAGRAPH)
    style.font.bold = True
    style.font.size = Pt(BODY_SIZE_PT)

    _heading(document.add_paragraph(), heading)
    paragraph = document.add_paragraph(style=BOLD_STYLE)
    paragraph.add_run(INHERITS_BOLD)
    paragraph.add_run(SWITCHED_OFF).font.bold = False

    split = document.add_paragraph()
    for part in SPLIT_RUNS:
        _text(split, part)

    document.save(str(path))
    return path


def _project_table(document: WordDocument, project: Project) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for column, width in zip(table.columns, (PERIOD_COLUMN_MM, DETAIL_COLUMN_MM), strict=True):
        column.width = Mm(width)
        for cell in column.cells:
            cell.width = Mm(width)

    period, detail = table.rows[0].cells
    _text(period.paragraphs[0], project.period, bold=True)
    _text(period.add_paragraph(), "Rolle:", bold=True)
    _text(period.add_paragraph(), project.role)
    # A blank line inside a cell is layout, not spacing between blocks.
    period.add_paragraph()
    _text(period.add_paragraph(), "Kunde:", bold=True)
    _text(period.add_paragraph(), project.customer)

    _text(detail.paragraphs[0], "Tätigkeiten", bold=True)
    for activity in project.activities:
        _bullet(detail.add_paragraph(), activity)
    for note in project.details:
        _bullet(detail.add_paragraph(), note, level=1)
    if project.link is not None:
        _link(detail.add_paragraph(), *project.link)


def _heading(paragraph: Paragraph, text: str) -> Paragraph:
    _text(paragraph, text, bold=True, size_pt=HEADING_SIZE_PT)
    return paragraph


def _text(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
    size_pt: float = BODY_SIZE_PT,
) -> Paragraph:
    run = paragraph.add_run(text)
    run.font.size = Pt(size_pt)
    if bold:
        run.font.bold = True
    return paragraph


def _bullet(paragraph: Paragraph, text: str, *, level: int = 0) -> Paragraph:
    """A bullet the way the real document makes them: numbering on the paragraph.

    Not with Word's ``List Bullet`` style, because the source file does not --
    it numbers a heading style, then switches the heading's formatting off run by
    run, which is exactly the case the importer has to resolve correctly.
    """
    _text(paragraph, text)
    numbering = OxmlElement("w:numPr")
    for tag, value in (("w:ilvl", level), ("w:numId", 1)):
        element = OxmlElement(tag)
        element.set(qn("w:val"), str(value))
        numbering.append(element)
    paragraph._p.get_or_add_pPr().append(numbering)
    return paragraph


def _link(paragraph: Paragraph, text: str, url: str) -> Paragraph:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = paragraph.add_run(text)
    run.font.size = Pt(BODY_SIZE_PT)
    hyperlink.append(run._r)
    paragraph._p.append(hyperlink)
    return paragraph


def cell_text(cell: _Cell) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs)
