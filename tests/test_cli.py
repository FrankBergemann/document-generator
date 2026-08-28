from __future__ import annotations

from pathlib import Path

import docx
import pytest

from cv_generator.cli import main
from cv_generator.pdf import BROWSER_WS_ENV
from tests.conftest import PROJECTS_CONFIG, PROJECTS_MD, write_config
from tests.support import (
    CHROME_AVAILABLE,
    LOCAL_CHROME_INSTALLED,
    PROJEKTLISTE_NAME,
    requires_chromium,
)


class TestBuildHTML:
    def test_writes_html(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        out = tmp_path / "document.html"
        assert main(["build", str(minimal_path), "-o", str(out)]) == 0
        assert "Ada Lovelace" in out.read_text(encoding="utf-8")
        assert str(out) in capsys.readouterr().out

    def test_creates_missing_output_directory(self, tmp_path: Path, minimal_path: Path) -> None:
        out = tmp_path / "nested" / "deeper" / "document.html"
        assert main(["build", str(minimal_path), "-o", str(out)]) == 0
        assert out.is_file()

    def test_default_output_path_follows_the_format(
        self, tmp_path: Path, minimal_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.chdir(tmp_path)
        assert main(["build", str(minimal_path), "-f", "html"]) == 0
        assert (tmp_path / "dist" / "minimal.html").is_file()

    def test_unknown_theme_is_a_clean_error(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        code = main(
            ["build", str(minimal_path), "-t", "nope", "-o", str(tmp_path / "document.html")]
        )
        assert code == 1
        assert "not found" in capsys.readouterr().err

    def test_missing_source_is_a_clean_error(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        assert main(["build", str(tmp_path / "nope.md")]) == 1
        assert "cannot read" in capsys.readouterr().err


class TestBuildEveryFormat:
    """No `-f` renders the whole set, so one run produces what gets sent out."""

    @requires_chromium
    def test_no_format_writes_all_of_them(
        self,
        tmp_path: Path,
        minimal_path: Path,
        monkeypatch: pytest.MonkeyPatch,
        capsys: pytest.CaptureFixture[str],
    ) -> None:
        monkeypatch.chdir(tmp_path)
        assert main(["build", str(minimal_path)]) == 0
        for suffix in ("html", "docx", "pdf"):
            assert (tmp_path / "dist" / f"minimal.{suffix}").is_file()
        # One line per format: run/build.sh archives each of them by parsing these.
        out = capsys.readouterr().out
        assert [line.split(".")[-1] for line in out.splitlines() if line.startswith("wrote ")] == [
            "html",
            "docx",
            "pdf",
        ]

    def test_format_is_repeatable_and_keeps_its_order(
        self,
        tmp_path: Path,
        minimal_path: Path,
        monkeypatch: pytest.MonkeyPatch,
        capsys: pytest.CaptureFixture[str],
    ) -> None:
        monkeypatch.chdir(tmp_path)
        argv = ["build", str(minimal_path), "-f", "docx", "-f", "html", "-f", "docx"]
        assert main(argv) == 0
        wrote = [line for line in capsys.readouterr().out.splitlines() if line.startswith("wrote ")]
        assert [Path(line.removeprefix("wrote ")).suffix for line in wrote] == [".docx", ".html"]

    def test_out_takes_the_format_from_its_extension(
        self, tmp_path: Path, minimal_path: Path
    ) -> None:
        out = tmp_path / "document.docx"
        assert main(["build", str(minimal_path), "-o", str(out)]) == 0
        assert "Ada Lovelace" in [p.text for p in docx.Document(str(out)).paragraphs]

    def test_out_with_an_unknown_extension_is_a_clean_error(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        assert main(["build", str(minimal_path), "-o", str(tmp_path / "document.rtf")]) == 1
        assert "cannot tell the format" in capsys.readouterr().err

    def test_a_failing_format_does_not_cost_the_others(
        self,
        tmp_path: Path,
        minimal_path: Path,
        monkeypatch: pytest.MonkeyPatch,
        capsys: pytest.CaptureFixture[str],
    ) -> None:
        # An unreachable PDF engine stands in for the everyday case, a machine
        # with no browser: the .html and .docx are still worth having, and the
        # exit code still says the run was not clean.
        monkeypatch.chdir(tmp_path)
        assert main(["build", str(minimal_path), "-e", "nope"]) == 1
        assert (tmp_path / "dist" / "minimal.html").is_file()
        assert (tmp_path / "dist" / "minimal.docx").is_file()
        assert not (tmp_path / "dist" / "minimal.pdf").exists()
        assert "unknown PDF engine" in capsys.readouterr().err

    def test_out_takes_a_single_format(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        out = tmp_path / "document.html"
        assert main(["build", str(minimal_path), "-f", "html", "-o", str(out)]) == 0
        assert main(["build", str(minimal_path), "-f", "html", "-f", "docx", "-o", str(out)]) == 1
        assert "takes a single --format" in capsys.readouterr().err


class TestBuildFromARecipe:
    """A `.json` source assembles the document; a `.md` source is the document."""

    def test_the_output_is_named_by_the_recipe_not_by_the_recipe_file(
        self,
        projects_path: Path,
        tmp_path_factory: pytest.TempPathFactory,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        # Without this every project would ship a dist/config.html.
        monkeypatch.chdir(tmp_path_factory.mktemp("cwd"))
        assert main(["build", str(projects_path), "-f", "html"]) == 0
        assert (Path("dist") / "document.html").is_file()

    def test_output_overrides_that_name(
        self,
        projects_dir: Path,
        tmp_path_factory: pytest.TempPathFactory,
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        source = write_config(projects_dir, PROJECTS_CONFIG | {"output": "lebenslauf"})
        monkeypatch.chdir(tmp_path_factory.mktemp("cwd"))
        assert main(["build", str(source), "-f", "html"]) == 0
        assert (Path("dist") / "lebenslauf.html").is_file()

    def test_every_source_reaches_the_output(self, tmp_path: Path, projects_path: Path) -> None:
        out = tmp_path / "document.html"
        assert main(["build", str(projects_path), "-o", str(out)]) == 0
        html = out.read_text(encoding="utf-8")
        assert "Land Schleswig-Holstein" in html  # from the .docx
        assert "Python" in html  # from the .md
        assert "darf nicht" not in html  # a span no entry asked for

    def test_out_still_wins(self, tmp_path: Path, projects_path: Path) -> None:
        out = tmp_path / "somewhere" / "other.html"
        assert main(["build", str(projects_path), "-o", str(out)]) == 0
        assert out.is_file()


class TestBuildWithPhoto:
    def test_html_embeds_the_photo(self, tmp_path: Path, photo_path: Path) -> None:
        out = tmp_path / "document.html"
        assert main(["build", str(photo_path), "-o", str(out)]) == 0
        assert 'src="data:image/png;base64,' in out.read_text(encoding="utf-8")

    def test_docx_embeds_the_photo(self, tmp_path: Path, photo_path: Path) -> None:
        out = tmp_path / "document.docx"
        assert main(["build", str(photo_path), "-f", "docx", "-o", str(out)]) == 0
        assert len(docx.Document(str(out)).inline_shapes) == 1


class TestBuildDOCX:
    def test_writes_a_readable_docx(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        out = tmp_path / "document.docx"
        assert main(["build", str(minimal_path), "-f", "docx", "-o", str(out)]) == 0
        assert "Ada Lovelace" in [p.text for p in docx.Document(str(out)).paragraphs]
        assert str(out) in capsys.readouterr().out

    def test_default_output_path(
        self, tmp_path: Path, minimal_path: Path, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.chdir(tmp_path)
        assert main(["build", str(minimal_path), "-f", "docx"]) == 0
        assert (tmp_path / "dist" / "minimal.docx").is_file()

    def test_needs_no_pdf_engine(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        # docx must work on a machine without Chromium, so it must not touch
        # the engine registry at all.
        code = main(
            [
                "build",
                str(minimal_path),
                "-f",
                "docx",
                "-e",
                "nonsense",
                "-o",
                str(tmp_path / "a.docx"),
            ]
        )
        assert code == 0
        assert capsys.readouterr().err == ""


class TestBuildPDF:
    @requires_chromium
    def test_writes_a_pdf(self, tmp_path: Path, minimal_path: Path) -> None:
        out = tmp_path / "document.pdf"
        assert main(["build", str(minimal_path), "-f", "pdf", "-o", str(out)]) == 0
        assert out.read_bytes().startswith(b"%PDF-")

    @pytest.mark.skipif(CHROME_AVAILABLE or LOCAL_CHROME_INSTALLED, reason="a browser is available")
    def test_without_chromium_it_explains_the_install(
        self,
        tmp_path: Path,
        minimal_path: Path,
        capsys: pytest.CaptureFixture[str],
        monkeypatch: pytest.MonkeyPatch,
    ) -> None:
        monkeypatch.delenv(BROWSER_WS_ENV, raising=False)
        out = tmp_path / "document.pdf"
        assert main(["build", str(minimal_path), "-f", "pdf", "-o", str(out)]) == 1
        assert "playwright install chromium" in capsys.readouterr().err
        assert not out.exists()

    def test_unknown_engine_is_a_clean_error(
        self, tmp_path: Path, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        code = main(
            [
                "build",
                str(minimal_path),
                "-f",
                "pdf",
                "-e",
                "nope",
                "-o",
                str(tmp_path / "document.pdf"),
            ]
        )
        assert code == 1
        assert "unknown PDF engine" in capsys.readouterr().err


class TestValidate:
    def test_reports_sections(self, minimal_path: Path, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["validate", str(minimal_path)]) == 0
        out = capsys.readouterr().out
        assert "ok" in out
        assert "Experience (experience)" in out

    def test_reports_the_photo_it_found(
        self, photo_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        # The one file the CV points at, so "was it found and read?" is worth
        # answering before a build.
        assert main(["validate", str(photo_path)]) == 0
        assert "photo: image/png" in capsys.readouterr().out

    def test_says_nothing_about_a_photo_without_one(
        self, minimal_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        assert main(["validate", str(minimal_path)]) == 0
        assert "photo" not in capsys.readouterr().out

    def test_unreadable_photo_is_a_clean_error(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        source = tmp_path / "document.md"
        source.write_text("---\nname: Ada\nphoto: gone.png\n---\n", encoding="utf-8")
        assert main(["validate", str(source)]) == 1
        assert "cannot read photo" in capsys.readouterr().err

    def test_reports_the_file_behind_each_section(
        self, projects_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        # A recipe resolves globs and headlines, and neither is visible in the
        # result, so which file fed which section is the thing to check before a
        # build goes out.
        assert main(["validate", str(projects_path)]) == 0
        out = capsys.readouterr().out
        assert "Projekte (projekte)" in out
        assert PROJEKTLISTE_NAME in out
        assert "block(s) from" in out
        assert "Kenntnisse (kenntnisse) <- " in out
        assert "document.md" in out

    def test_reports_the_name_the_outputs_will_take(
        self, projects_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        # The recipe is called config.json and the result is not.
        assert main(["validate", str(projects_path)]) == 0
        assert "-> document.*" in capsys.readouterr().out

    def test_a_broken_recipe_is_a_clean_error(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        (tmp_path / "document.md").write_text(PROJECTS_MD, encoding="utf-8")
        source = write_config(
            tmp_path,
            {
                "sections": [
                    {"source": "document.md", "end": "Berufserfahrung"},
                    {"source": "*Projektliste*.docx", "begin": "Projekthistorie"},
                ]
            },
        )
        assert main(["validate", str(source)]) == 1
        assert "Projektliste" in capsys.readouterr().err

    def test_invalid_json_is_a_clean_error(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        bad = tmp_path / "config.json"
        bad.write_text("{not json", encoding="utf-8")
        assert main(["validate", str(bad)]) == 1
        assert "invalid JSON" in capsys.readouterr().err

    def test_invalid_file_exits_nonzero(
        self, tmp_path: Path, capsys: pytest.CaptureFixture[str]
    ) -> None:
        bad = tmp_path / "bad.md"
        bad.write_text("no frontmatter here\n", encoding="utf-8")
        assert main(["validate", str(bad)]) == 1
        assert "missing YAML frontmatter" in capsys.readouterr().err


class TestInfoCommands:
    def test_engines_reports_chrome_as_implemented(
        self, capsys: pytest.CaptureFixture[str]
    ) -> None:
        assert main(["engines"]) == 0
        out = capsys.readouterr().out
        assert "implemented: chrome" in out
        assert "weasyprint [not implemented]" in out

    @requires_chromium
    def test_engines_reports_chrome_as_ready(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["engines"]) == 0
        assert "chrome [ready]" in capsys.readouterr().out

    def test_engines_names_the_browser_server_when_one_is_configured(
        self, capsys: pytest.CaptureFixture[str], monkeypatch: pytest.MonkeyPatch
    ) -> None:
        # Otherwise a stopped browser container is indistinguishable from a
        # missing local install.
        monkeypatch.setenv(BROWSER_WS_ENV, "ws://playwright:3000/")
        assert main(["engines"]) == 0
        assert "browser server: ws://playwright:3000/" in capsys.readouterr().out

    def test_engines_says_nothing_about_a_server_without_one(
        self, capsys: pytest.CaptureFixture[str], monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.delenv(BROWSER_WS_ENV, raising=False)
        assert main(["engines"]) == 0
        assert "browser server" not in capsys.readouterr().out

    def test_themes_lists_classic_and_notes_docx(self, capsys: pytest.CaptureFixture[str]) -> None:
        assert main(["themes"]) == 0
        out = capsys.readouterr().out
        assert "classic" in out
        assert "docx" in out


class TestArgumentParsing:
    def test_subcommand_is_required(self) -> None:
        with pytest.raises(SystemExit) as excinfo:
            main([])
        assert excinfo.value.code == 2

    def test_unknown_format_is_rejected(self) -> None:
        with pytest.raises(SystemExit) as excinfo:
            main(["build", "-f", "rtf"])
        assert excinfo.value.code == 2

    def test_version(self, capsys: pytest.CaptureFixture[str]) -> None:
        with pytest.raises(SystemExit) as excinfo:
            main(["--version"])
        assert excinfo.value.code == 0
        assert "0.2.0" in capsys.readouterr().out
