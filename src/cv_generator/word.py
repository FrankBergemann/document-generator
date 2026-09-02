"""MS Word (.docx) output.

Word has no stylesheet, so the visual decisions that ``templates/classic/document.css``
expresses in CSS are mirrored here as a :class:`WordTheme` and applied as named
paragraph styles -- which also means the recipient can restyle the whole
document from Word's style pane instead of reformatting paragraph by paragraph.

Content comes from the same Markdown syntax tree the HTML renderer uses, so the
two formats cannot drift: a Markdown feature is either supported by both or
visibly missing from both. A section imported from an existing Word document
(see :mod:`cv_generator.docx_import`) arrives as blocks instead, and is written
with the formatting it had there rather than with this theme's styles.
"""

from __future__ import annotations

import io
from collections.abc import Sequence
from dataclasses import dataclass, replace
from pathlib import Path
from typing import Protocol

import docx
from docx.document import Document as WordDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Mm, Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from markdown_it.tree import SyntaxTreeNode

from cv_generator.errors import RenderError
from cv_generator.markdown import inline_children, to_tree
from cv_generator.models import (
    Contact,
    Document,
    Photo,
    RichBlock,
    RichParagraph,
    RichRun,
    RichTable,
)
from cv_generator.ooxml import (
    add_bottom_border,
    add_table_bottom_border,
    move_run_into,
    remove_paragraph,
    set_column_widths,
    set_language,
    set_table_cell_margins,
    start_hyperlink,
)

STYLE_NAME = "Document Name"
STYLE_HEADLINE = "Document Headline"
STYLE_CONTACT = "Document Contact"
STYLE_SUMMARY = "Document Summary"
STYLE_SECTION = "Document Section Heading"
STYLE_ENTRY = "Document Entry"
STYLE_META = "Document Entry Meta"
STYLE_BODY = "Document Body"
STYLE_CODE = "Document Code"
STYLE_QUOTE = "Document Quote"

CONTACT_SEPARATOR = "   ·   "

A4_WIDTH_MM = 210.0
A4_HEIGHT_MM = 297.0


class _Blocks(Protocol):
    """What the header/footer builders need of a document, table cell, or
    page header/footer part.

    With a photo the CV header's text goes into a table cell instead of
    straight into the document body; a page header/footer is a part of its
    own. All three offer ``add_paragraph`` with this signature; ``paragraphs``
    is used only to find the blank placeholder a fresh page header/footer
    part starts with (see ``WordRenderer._add_page_part``).
    """

    def add_paragraph(self, text: str = "", style: str | None = None) -> Paragraph: ...

    @property
    def paragraphs(self) -> list[Paragraph]: ...


@dataclass(frozen=True)
class WordTheme:
    """Typography and colour for ``.docx`` output.

    Defaults mirror the ``classic`` HTML theme. The body font is Calibri rather
    than the CSS stack's Source Sans 3 because a ``.docx`` has to render on a
    machine that has never heard of this project.
    """

    body_font: str = "Calibri"
    mono_font: str = "Consolas"

    body_size: float = 10.5
    name_size: float = 20.0
    headline_size: float = 12.0
    contact_size: float = 9.0
    section_size: float = 10.0
    entry_size: float = 11.0
    meta_size: float = 9.5

    ink: str = "1A1A1A"
    ink_muted: str = "55585C"
    accent: str = "2F5D8A"
    rule: str = "D4D7DB"

    page_margin_vertical_mm: float = 16.0
    page_margin_horizontal_mm: float = 15.0

    # Mirrors `--photo-width` / `--photo-max-height` in document.css: the photo sets
    # the header's height rather than being squeezed into it.
    photo_width_mm: float = 34.0
    # Only a guard: clears a 3:4 portrait (45.3mm at this width) untouched.
    photo_max_height_mm: float = 48.0
    photo_gap_mm: float = 8.0

    # Cell insets for a table imported from another document, in points. The
    # source's own value is not carried over: Word leaves most of these unset
    # and lets its default template decide, so reproducing it means reproducing
    # that template. Mirrors `.cv-block-table td { padding }` in document.css.
    imported_cell_padding_vertical_pt: float = 2.8
    imported_cell_padding_horizontal_pt: float = 4.0

    @property
    def content_width_mm(self) -> float:
        return A4_WIDTH_MM - 2 * self.page_margin_horizontal_mm


@dataclass(frozen=True)
class _Fmt:
    """Inline formatting accumulated while descending the syntax tree."""

    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False
    link: bool = False


_PLAIN = _Fmt()


class WordRenderer:
    """Writes a Document as a ``.docx`` file."""

    def __init__(self, theme: WordTheme | None = None) -> None:
        self.theme = theme or WordTheme()

    def render(self, doc: Document, output: Path) -> None:
        """Write ``doc`` to ``output`` as a Word document.

        Raises:
            RenderError: if the file cannot be written.
        """
        document = self._new_document(doc)
        if doc.has_identity():
            self._add_header(document, doc)

        if doc.summary:
            self._add_blocks(document, to_tree(doc.summary).children, STYLE_SUMMARY)

        for section in doc.sections:
            if section.title:
                self._add_section_heading(document, section.title)
            if section.blocks:
                self._add_imported(document, document, section.blocks)
            else:
                self._add_blocks(document, to_tree(section.markdown).children, STYLE_BODY)

        # Real, repeating page header/footer here, unlike HTML/PDF (see
        # document.html.j2): Word shows each on every page on its own.
        if doc.page_header:
            self._add_page_part(document, document.sections[-1].header, doc.page_header)
        if doc.page_footer:
            self._add_page_part(document, document.sections[-1].footer, doc.page_footer)

        output.parent.mkdir(parents=True, exist_ok=True)
        try:
            document.save(str(output))
        except OSError as exc:
            raise RenderError(f"cannot write {output}: {exc}") from exc

    def _add_page_part(
        self, document: WordDocument, part: _Blocks, blocks: Sequence[RichBlock]
    ) -> None:
        """Write ``blocks`` into a freshly defined page header or footer.

        Accessing ``.paragraphs`` is what creates the part in the first place,
        and Word pre-populates a new one with a single blank paragraph -- the
        same placeholder pattern as the photo header's table cell.
        """
        placeholder = part.paragraphs[0]
        self._add_imported(document, part, blocks)
        remove_paragraph(placeholder)

    # -- document setup ---------------------------------------------------

    def _new_document(self, doc: Document) -> WordDocument:
        document = docx.Document()
        self._define_styles(document, doc.lang)

        page = document.sections[0]
        page.page_width = Mm(A4_WIDTH_MM)
        page.page_height = Mm(A4_HEIGHT_MM)
        page.top_margin = page.bottom_margin = Mm(self.theme.page_margin_vertical_mm)
        page.left_margin = page.right_margin = Mm(self.theme.page_margin_horizontal_mm)

        document.core_properties.title = doc.name or ""
        document.core_properties.author = doc.name or ""
        return document

    def _define_styles(self, document: WordDocument, lang: str) -> None:
        theme = self.theme
        normal = document.styles["Normal"]
        normal.font.name = theme.body_font
        normal.font.size = Pt(theme.body_size)
        normal.font.color.rgb = RGBColor.from_string(theme.ink)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.line_spacing = 1.15
        set_language(normal, lang)

        def define(
            name: str,
            *,
            size: float,
            bold: bool = False,
            italic: bool = False,
            all_caps: bool = False,
            color: str | None = None,
            font: str | None = None,
            space_before: float = 0.0,
            space_after: float = 0.0,
            keep_with_next: bool = False,
        ) -> None:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
            style.font.size = Pt(size)
            style.font.bold = bold
            style.font.italic = italic
            style.font.all_caps = all_caps
            if font is not None:
                style.font.name = font
            if color is not None:
                style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(space_before)
            style.paragraph_format.space_after = Pt(space_after)
            style.paragraph_format.keep_with_next = keep_with_next

        define(STYLE_NAME, size=theme.name_size, bold=True, space_after=1, keep_with_next=True)
        define(
            STYLE_HEADLINE,
            size=theme.headline_size,
            color=theme.accent,
            space_after=3,
            keep_with_next=True,
        )
        define(STYLE_CONTACT, size=theme.contact_size, color=theme.ink_muted, space_after=2)
        define(
            STYLE_SUMMARY,
            size=theme.body_size,
            color=theme.ink_muted,
            space_before=8,
            space_after=3,
        )
        # Mirrors `.cv-section > h2`: small, bold, uppercase, accent, ruled.
        define(
            STYLE_SECTION,
            size=theme.section_size,
            bold=True,
            all_caps=True,
            color=theme.accent,
            space_before=12,
            space_after=4,
            keep_with_next=True,
        )
        define(STYLE_ENTRY, size=theme.entry_size, bold=True, space_before=7, keep_with_next=True)
        define(STYLE_META, size=theme.meta_size, color=theme.ink_muted, space_after=2)
        define(STYLE_BODY, size=theme.body_size, space_after=3)
        define(STYLE_CODE, size=theme.body_size - 1, font=theme.mono_font, space_after=3)
        define(STYLE_QUOTE, size=theme.body_size, italic=True, color=theme.ink_muted, space_after=3)

    # -- header -----------------------------------------------------------

    def _add_header(self, document: WordDocument, doc: Document) -> None:
        if doc.photo is None:
            last = self._add_identity(document, doc)
            # `.cv-header { border-bottom: 1.5px solid var(--accent) }`
            add_bottom_border(last, color=self.theme.accent, size=12, space=6)
            return

        # Word has no flexbox, so the two-column header is a borderless layout
        # table: identity on the left, photo hard against the right margin.
        theme = self.theme
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        photo_column_mm = theme.photo_width_mm + theme.photo_gap_mm
        set_column_widths(
            table,
            Mm(theme.content_width_mm - photo_column_mm),
            Mm(photo_column_mm),
        )
        # Zeroing the insets (the default) keeps the name flush with the body
        # text; the bottom one is the gap the rule needs, standing in for
        # `.cv-header { padding-bottom }`.
        set_table_cell_margins(table, bottom=6)

        identity_cell, photo_cell = table.rows[0].cells
        placeholder = identity_cell.paragraphs[0]
        self._add_identity(identity_cell, doc)
        remove_paragraph(placeholder)
        self._add_photo(photo_cell, doc.photo)

        add_table_bottom_border(table, color=theme.accent, size=12)

    def _add_identity(self, blocks: _Blocks, doc: Document) -> Paragraph:
        """Write name, headline and contact line; returns the last paragraph.

        ``doc.name`` may be absent -- a document with no Markdown source has no
        frontmatter to take it from -- so the paragraph is written empty rather
        than skipped, which keeps this method's "add exactly one paragraph per
        call" shape and so the header's own bottom rule (drawn on the paragraph
        this method returns) always has one to draw on.
        """
        last = blocks.add_paragraph(doc.name or "", style=STYLE_NAME)
        if doc.headline:
            last = blocks.add_paragraph(doc.headline, style=STYLE_HEADLINE)
        if not doc.contact.is_empty():
            last = self._add_contact(blocks, doc.contact)
        return last

    def _add_photo(self, cell: _Cell, photo: Photo) -> None:
        """Place the photo in ``cell``, right-aligned and scaled like the CSS."""
        paragraph = cell.paragraphs[0]
        paragraph.style = STYLE_BODY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        theme = self.theme
        picture = paragraph.add_run().add_picture(
            io.BytesIO(photo.data), width=Mm(theme.photo_width_mm)
        )
        # `max-height` in the stylesheet, applied by hand: scale both dimensions
        # so an unusually tall portrait is bounded without being distorted.
        limit = Mm(theme.photo_max_height_mm)
        if picture.height > limit:
            picture.width = Emu(round(picture.width * limit / picture.height))
            picture.height = limit

    def _add_contact(self, document: _Blocks, contact: Contact) -> Paragraph:
        paragraph = document.add_paragraph(style=STYLE_CONTACT)
        parts: list[tuple[str | None, str]] = []
        if contact.email:
            parts.append((f"mailto:{contact.email}", contact.email))
        if contact.phone:
            parts.append((None, contact.phone))
        if contact.location:
            parts.append((None, contact.location))
        parts.extend((link.url, link.label) for link in contact.links)

        for index, (url, text) in enumerate(parts):
            if index:
                self._add_run(paragraph, CONTACT_SEPARATOR, _PLAIN)
            if url is None:
                self._add_run(paragraph, text, _PLAIN)
            else:
                self._add_run(paragraph, text, _Fmt(link=True), start_hyperlink(paragraph, url))
        return paragraph

    def _add_section_heading(self, document: WordDocument, title: str) -> None:
        paragraph = document.add_paragraph(title, style=STYLE_SECTION)
        add_bottom_border(paragraph, color=self.theme.rule)

    # -- block content ----------------------------------------------------

    def _add_blocks(
        self, document: WordDocument, nodes: list[SyntaxTreeNode], default_style: str
    ) -> None:
        for index, node in enumerate(nodes):
            previous = nodes[index - 1] if index else None
            self._add_block(document, node, previous, default_style)

    def _add_block(
        self,
        document: WordDocument,
        node: SyntaxTreeNode,
        previous: SyntaxTreeNode | None,
        default_style: str,
    ) -> None:
        kind = node.type

        if kind == "heading":
            paragraph = document.add_paragraph(style=STYLE_ENTRY)
            self._add_inline(paragraph, inline_children(node))

        elif kind == "paragraph":
            style = STYLE_META if _is_entry_meta(node, previous) else default_style
            paragraph = document.add_paragraph(style=style)
            self._add_inline(paragraph, inline_children(node))

        elif kind in ("bullet_list", "ordered_list"):
            self._add_list(document, node, depth=1)

        elif kind in ("fence", "code_block"):
            self._add_code(document, node.content)

        elif kind == "blockquote":
            self._add_blocks(document, node.children, STYLE_QUOTE)

        elif kind == "hr":
            add_bottom_border(document.add_paragraph(style=STYLE_BODY), color=self.theme.rule)

        elif kind == "table":
            self._add_table(document, node)

        elif node.children:
            # Unknown container (e.g. a plugin block): keep its content rather
            # than dropping it silently.
            self._add_blocks(document, node.children, default_style)

    def _add_list(self, document: WordDocument, node: SyntaxTreeNode, *, depth: int) -> None:
        ordered = node.type == "ordered_list"
        style = _list_style(document, ordered=ordered, depth=depth)

        for item in node.children:
            for child in item.children:
                if child.type in ("bullet_list", "ordered_list"):
                    self._add_list(document, child, depth=depth + 1)
                elif child.type == "paragraph":
                    paragraph = document.add_paragraph(style=style)
                    paragraph.paragraph_format.space_after = Pt(2)
                    self._add_inline(paragraph, inline_children(child))
                else:
                    self._add_block(document, child, None, STYLE_BODY)

    def _add_code(self, document: WordDocument, content: str) -> None:
        paragraph = document.add_paragraph(style=STYLE_CODE)
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run()
        for index, line in enumerate(content.rstrip("\n").split("\n")):
            if index:
                run.add_break()
            run.add_text(line)

    def _add_table(self, document: WordDocument, node: SyntaxTreeNode) -> None:
        header_rows, body_rows = _table_rows(node)
        rows = header_rows + body_rows
        if not rows:
            return

        table = document.add_table(rows=0, cols=max(len(row) for row in rows))
        if "Table Grid" in _style_names(document):
            table.style = "Table Grid"

        for index, cells in enumerate(rows):
            row = table.add_row()
            is_header = index < len(header_rows)
            for column, cell in enumerate(cells):
                paragraph = row.cells[column].paragraphs[0]
                paragraph.style = document.styles[STYLE_BODY]
                self._add_inline(paragraph, inline_children(cell), _Fmt(bold=is_header))

    # -- imported blocks --------------------------------------------------

    def _add_imported(
        self, document: WordDocument, target: _Blocks, blocks: Sequence[RichBlock]
    ) -> None:
        """Write blocks imported from another Word document.

        ``target`` is where the paragraphs go -- the document, or a table cell --
        while ``document`` stays the document, because a cell has no style
        collection of its own to look list styles up in.

        The blank paragraph the source kept between two tables is dropped on
        import (see :mod:`cv_generator.docx_import`), but Word still needs
        *something* between two adjacent ``w:tbl`` elements or it renders them
        as one merged table -- so a table that directly follows another one
        here gets an empty paragraph reinserted ahead of it.
        """
        previous_was_table = False
        for block in blocks:
            if isinstance(block, RichTable):
                if previous_was_table:
                    target.add_paragraph()
                self._add_imported_table(document, block)
            else:
                self._add_imported_paragraph(document, target, block)
            previous_was_table = isinstance(block, RichTable)

    def _add_imported_paragraph(
        self, document: WordDocument, target: _Blocks, block: RichParagraph
    ) -> None:
        style = STYLE_BODY
        if block.level is not None:
            style = _list_style(document, ordered=block.ordered, depth=block.level + 1)
        paragraph = target.add_paragraph(style=style)
        for run in block.runs:
            self._add_imported_run(paragraph, run)

    def _add_imported_run(self, paragraph: Paragraph, run: RichRun) -> None:
        """Write one run with the formatting it had in the source document.

        Still only ever switches attributes *on* (see ``_add_run``): the styles
        these paragraphs use carry none of them, so "not set in the source" and
        "off" produce the same document -- and `bold = False` would fight the
        style rather than the source.
        """
        link = start_hyperlink(paragraph, run.link) if run.link else None
        word_run = paragraph.add_run(run.text)
        if run.image is not None:
            # No text formatting applies to a picture. The source's own size
            # (`wp:extent`) is not carried over, only the image itself -- Word
            # picks a default size the same way it would for a pasted image.
            word_run.add_picture(io.BytesIO(run.image.data))
        else:
            font = word_run.font
            if run.bold:
                font.bold = True
            if run.italic:
                font.italic = True
            if run.underline:
                font.underline = True
            if run.strike:
                font.strike = True
            if run.size_pt is not None:
                font.size = Pt(run.size_pt)
            if run.color is not None:
                font.color.rgb = RGBColor.from_string(run.color)
            elif run.link:
                font.color.rgb = RGBColor.from_string(self.theme.accent)
        if link is not None:
            move_run_into(word_run, link)

    def _add_imported_table(self, document: WordDocument, block: RichTable) -> None:
        columns = max((len(row) for row in block.rows), default=0)
        if not columns:
            return

        theme = self.theme
        table = document.add_table(rows=0, cols=columns)
        if block.bordered and "Table Grid" in _style_names(document):
            table.style = "Table Grid"
        if block.centered:
            # Sized to its own content instead of stretched to the page -- the
            # default a table gets when no column widths are set at all -- and
            # centered between the margins, the way Word centers a table left
            # otherwise sitting flush against the left margin.
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
        elif len(block.column_widths) == columns:
            set_column_widths(
                table, *(Mm(theme.content_width_mm * share) for share in block.column_widths)
            )
        set_table_cell_margins(
            table,
            top=theme.imported_cell_padding_vertical_pt,
            left=theme.imported_cell_padding_horizontal_pt,
            bottom=theme.imported_cell_padding_vertical_pt,
            right=theme.imported_cell_padding_horizontal_pt,
        )

        for row in block.rows:
            cells = table.add_row().cells
            for index, cell in enumerate(row):
                if not cell.paragraphs:
                    # Word repairs a cell with no paragraph in it, so an empty
                    # cell keeps the one python-docx put there.
                    continue
                placeholder = cells[index].paragraphs[0]
                self._add_imported(document, cells[index], cell.paragraphs)
                remove_paragraph(placeholder)

    # -- inline content ---------------------------------------------------

    def _add_inline(
        self,
        paragraph: Paragraph,
        nodes: list[SyntaxTreeNode],
        fmt: _Fmt = _PLAIN,
        container: object | None = None,
    ) -> None:
        for node in nodes:
            kind = node.type

            if kind == "text":
                self._add_run(paragraph, node.content, fmt, container)
            elif kind == "code_inline":
                self._add_run(paragraph, node.content, replace(fmt, code=True), container)
            elif kind == "strong":
                self._add_inline(paragraph, node.children, replace(fmt, bold=True), container)
            elif kind == "em":
                self._add_inline(paragraph, node.children, replace(fmt, italic=True), container)
            elif kind == "s":
                self._add_inline(paragraph, node.children, replace(fmt, strike=True), container)
            elif kind == "link":
                url = str(node.attrs.get("href", ""))
                link = start_hyperlink(paragraph, url) if url else container
                self._add_inline(paragraph, node.children, replace(fmt, link=bool(url)), link)
            elif kind == "softbreak":
                self._add_run(paragraph, " ", fmt, container)
            elif kind == "hardbreak":
                run = paragraph.add_run()
                run.add_break()
                if container is not None:
                    move_run_into(run, container)  # type: ignore[arg-type]
            elif kind == "image":
                # No image support yet; the alt text is better than nothing.
                if node.content:
                    self._add_run(paragraph, node.content, replace(fmt, italic=True), container)
            elif node.children:
                self._add_inline(paragraph, node.children, fmt, container)

    def _add_run(
        self,
        paragraph: Paragraph,
        text: str,
        fmt: _Fmt,
        container: object | None = None,
    ) -> None:
        run = paragraph.add_run(text)
        # Only ever turn attributes *on*: setting them False would override the
        # paragraph style, un-bolding entry headings and section titles.
        if fmt.bold:
            run.font.bold = True
        if fmt.italic:
            run.font.italic = True
        if fmt.strike:
            run.font.strike = True
        if fmt.code:
            run.font.name = self.theme.mono_font
        if fmt.link:
            run.font.color.rgb = RGBColor.from_string(self.theme.accent)
            run.font.underline = True
        if container is not None:
            move_run_into(run, container)  # type: ignore[arg-type]


def _is_entry_meta(node: SyntaxTreeNode, previous: SyntaxTreeNode | None) -> bool:
    """Whether a paragraph is an entry's date/place line.

    Mirrors the stylesheet's ``h3 + p > em:only-child``: a wholly italic
    paragraph directly beneath an entry heading.
    """
    if previous is None or previous.type != "heading":
        return False
    children = inline_children(node)
    return len(children) == 1 and children[0].type == "em"


def _table_rows(
    node: SyntaxTreeNode,
) -> tuple[list[list[SyntaxTreeNode]], list[list[SyntaxTreeNode]]]:
    """Split a table node into its header rows and its body rows."""
    header: list[list[SyntaxTreeNode]] = []
    body: list[list[SyntaxTreeNode]] = []
    for part in node.children:
        target = header if part.type == "thead" else body
        target.extend(list(row.children) for row in part.children)
    return header, body


def _style_names(document: WordDocument) -> set[str]:
    return {style.name for style in document.styles if style.name}


def _list_style(document: WordDocument, *, ordered: bool, depth: int) -> str:
    """Pick a built-in Word list style, degrading gracefully.

    Word's default template defines ``List Bullet`` through ``List Bullet 3``;
    deeper nesting reuses the deepest available style rather than failing.
    """
    base = "List Number" if ordered else "List Bullet"
    names = _style_names(document)
    for candidate in (f"{base} {depth}", base) if depth > 1 else (base,):
        if candidate in names:
            return candidate
    return STYLE_BODY
