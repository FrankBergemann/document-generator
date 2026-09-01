"""Read one section out of an existing ``.docx`` and into the Document model.

Some content is maintained in Word and nowhere else -- a project list with one
table per project, kept up to date because that is the file that gets sent
around. Retyping it as Markdown would lose both the layout (two columns, a
bullet list inside a cell) and the single source of truth, so instead the
section is imported: located by its heading, converted to the
:class:`~cv_generator.models.RichBlock` tree, and rendered by the normal
backends.

What "keeping the formatting" means here, precisely:

* **Carried over** -- bold, italic, underline, strikethrough, font size, colour,
  hyperlinks, bullet/number nesting, table columns and their widths, and whether
  the table is ruled.
* **Left to the Document's theme** -- font family, page geometry, paragraph spacing and
  the section heading itself. The imported blocks have to sit in a Document rendered
  by this project's themes, not carry a second document's page design into it.

Formatting is *resolved* rather than copied: a value set directly on the run
wins, otherwise the character style, otherwise the paragraph style, following
each ``basedOn`` chain. Word documents lean on both -- the list items in the
sample file use a heading style and then switch bold off run by run -- so
reading only one of the two gets the answer wrong.

Where a section ends is either stated or worked out. The build recipe may name
the heading the import stops before, which is the unambiguous answer; without one
the end is decided by *formatting*, because a hand-made Word Document usually has no
heading *styles* at all, only bold, larger text: the section then runs from its
heading to the next top-level paragraph that looks exactly like that heading
(same weight, same size), or to the end of the document. ``heading`` and ``end``
are regular expressions, matched case-insensitively and in full against a
paragraph's text -- plain heading text has no regex syntax in it, so it keeps
matching a paragraph that is exactly that text, the way an exact comparison did
before regular expressions were allowed. A full match, not a partial one: a
recipe naming "Rechnung" must not also stop at a paragraph that merely mentions
the word. A paragraph whose exact wording varies -- a closing line naming a
payment term, say -- can still be matched by covering the *whole* line with a
pattern, wildcards standing in for the part that changes.

Blank paragraphs are kept, not dropped: a source document's own spacing is
part of what it looked like, and this module's job is to keep formatting, not
to guess which whitespace was meaningful.

The private-API reads (``paragraph._p``, ``table._tbl``, ``cell._tc``,
``header._element``, ``footer._element``) stay here rather than moving to
:mod:`cv_generator.ooxml`. That module exists because *writing* OOXML has an
element-order trap; reading has none, and the alternative would be a dozen
one-line accessors in a module about the renderer's needs.
"""

from __future__ import annotations

import re
import zipfile
from collections.abc import Callable, Iterator, Sequence
from dataclasses import dataclass
from itertools import chain
from pathlib import Path
from typing import TypeVar

import docx
from docx.document import Document as WordDocument
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.section import _Footer, _Header
from docx.styles.style import CharacterStyle
from docx.table import Table as WordTable
from docx.text.font import Font
from docx.text.paragraph import Paragraph as WordParagraph
from docx.text.run import Run
from docx.types import ProvidesStoryPart

from cv_generator.errors import DocParseError
from cv_generator.models import RichBlock, RichCell, RichParagraph, RichRun, RichTable

# Word's lock file for an open document: same suffix, not a document. Without
# this, simply having the file open would turn discovery into "more than one".
LOCK_PREFIX = "~$"

# `qn()` only knows the namespace prefixes python-docx itself uses; markup
# compatibility (`mc:`, an OOXML extensibility mechanism, not one of them) is
# only needed here, for `_footer_blocks`.
_MC_FALLBACK = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"

_Block = WordParagraph | WordTable
_T = TypeVar("_T")


def load_section(path: Path, heading: str | None = None, end: str | None = None) -> list[RichBlock]:
    """Import the blocks under ``heading`` from the Word document at ``path``.

    Args:
        heading: A regular expression, matched case-insensitively and in full
            against a top-level paragraph's text (``re.fullmatch``) -- plain
            heading text, which has no regex syntax in it, matches itself
            exactly. A partial match would be too loose to locate a heading by:
            a title such as "Rechnung" would then also match a paragraph that
            merely mentions the word. ``None`` starts at the first block, which
            is the whole document when ``end`` is absent too.
        end: The heading the import stops *before*, matched the same way. Without
            it the end is found by formatting instead -- see :func:`_section_of`.

    Returns:
        The section's paragraphs and tables, in document order, blank
        paragraphs kept -- the source document's own spacing is content too,
        not an artifact to discard.

    Raises:
        DocParseError: if the file cannot be read as a Word document, ``heading``
            or ``end`` is not a valid regular expression, no paragraph matches,
            or there is nothing under it.
    """
    document = _open(path)
    blocks = list(_iter_blocks(document.element.body, document))
    numbering = _Numbering(document)

    section_blocks = _section_of(blocks, heading, end, path)
    if not section_blocks:
        raise DocParseError(f"{path}: nothing follows the heading {heading!r}")
    return [_convert(block, numbering) for block in section_blocks]


def load_footer(path: Path) -> list[RichBlock]:
    """The last section's footer, from the Word document at ``path``.

    Word may define up to three footers per section (default, first page, even
    page); which ones are used depends on the document's own page-numbering
    settings. The first page footer is tried first and the default one second,
    since a short, single-page document -- an invoice, typically -- shows only
    its first-page footer, which is then the one carrying real content while
    the "default" one (meant for a second page that never comes) is empty.

    See :func:`_hdrftr_blocks` for how the paragraphs themselves are read.

    Returns:
        The footer's paragraphs, in document order, blank ones kept. Empty if
        the document has no footer worth carrying over -- this is always
        optional, never an error.
    """
    document = _open(path)
    numbering = _Numbering(document)
    section = document.sections[-1]
    for footer in (section.first_page_footer, section.footer):
        blocks = _hdrftr_blocks(footer, numbering)
        if _has_text(blocks):
            return blocks
    return []


def load_header(path: Path) -> list[RichBlock]:
    """The last section's header, from the Word document at ``path``.

    The header counterpart of :func:`load_footer` -- same three-kind fallback
    (first page tried before default), same reading of the paragraphs
    (:func:`_hdrftr_blocks`), same "empty means no header, not an error".
    """
    document = _open(path)
    numbering = _Numbering(document)
    section = document.sections[-1]
    for header in (section.first_page_header, section.header):
        blocks = _hdrftr_blocks(header, numbering)
        if _has_text(blocks):
            return blocks
    return []


def _has_text(blocks: list[RichBlock]) -> bool:
    """Whether any block carries real text, as opposed to only blank paragraphs.

    Merely accessing a header/footer's ``.paragraphs`` (which
    :func:`_hdrftr_blocks` does) can create a fresh, empty definition with one
    blank placeholder paragraph of its own -- a side effect of python-docx's
    "read creates it if absent" behaviour, not a sign that this is the part
    with real content. Blank paragraphs are otherwise kept (see
    :func:`_hdrftr_blocks`), so this is not the same check as "no blocks at
    all"; it decides which of several candidate parts to read from, before any
    blank paragraph in the chosen one is preserved.
    """
    return any(isinstance(block, RichParagraph) and block.text() for block in blocks)


def _hdrftr_blocks(part: _Header | _Footer, numbering: _Numbering) -> list[RichBlock]:
    """A header or footer part's paragraphs, at any depth, each counted once.

    A letterhead-style header or footer is usually an anchored text box, not a
    plain paragraph -- its real content sits several levels deeper in the XML,
    inside ``w:txbxContent``, rather than as a direct child of the part. So
    every paragraph anywhere under it is read, not only top-level ones; the
    text box's own floating position is meaningless outside Word anyway, so
    flattening it into normal reading order is the best a page-agnostic format
    can do. Tables are not supported here (unlike a section import) -- a
    header or footer's own paragraphs are the common case, and nothing in this
    project's sample documents needs more.

    A text box commonly carries two equivalent copies of its content --
    ``mc:Choice`` (the DrawingML shape Word itself writes) and ``mc:Fallback``
    (a VML shape, for a Word old enough not to read the first) -- exactly one
    of which a real renderer would show. Reading both would double every
    paragraph inside one, so a paragraph under a ``mc:Fallback`` is skipped.

    Blank paragraphs are kept, same as a section import -- the source
    document's own spacing is content too, not an artifact to discard.
    """
    root = part._element
    paragraphs = (
        WordParagraph(element, part)
        for element in root.findall(f".//{qn('w:p')}")
        if not any(ancestor.tag == _MC_FALLBACK for ancestor in element.iterancestors())
    )
    return [_convert(p, numbering) for p in paragraphs]


# -- locating the section -------------------------------------------------


@dataclass(frozen=True)
class _HeadingShape:
    """How this document formats a section heading, learnt from the one we found."""

    bold: bool
    size_pt: float | None


def _open(path: Path) -> WordDocument:
    try:
        return docx.Document(str(path))
    except (PackageNotFoundError, zipfile.BadZipFile, KeyError, ValueError, OSError) as exc:
        raise DocParseError(f"cannot read {path} as a Word document: {exc}") from exc


def _iter_blocks(container: BaseOxmlElement, parent: ProvidesStoryPart) -> Iterator[_Block]:
    """Yield the paragraphs and tables of ``container`` in document order.

    python-docx exposes ``.paragraphs`` and ``.tables`` separately, which loses
    the order the two are interleaved in -- and here that order *is* the content.
    """
    for child in container.iterchildren():
        if child.tag == qn("w:p"):
            yield WordParagraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield WordTable(child, parent)


def _section_of(
    blocks: Sequence[_Block], heading: str | None, end: str | None, path: Path
) -> list[_Block]:
    """The blocks between ``heading`` and where the section ends.

    Named or guessed, in that order. A named ``end`` is a regular expression
    matched against another heading's text and says exactly where to stop, which
    is what a build recipe can state. Without one the end is found by
    *formatting*: a hand-made Word Document has no ``Heading 1`` anywhere -- its
    headings are bold, 12pt, no style at all -- so the import runs to the next
    top-level paragraph shaped like the one it started from. With no ``heading``
    to learn a shape from there is nothing to guess with, so the import simply
    runs to the end of the document.
    """
    # -1 rather than 0, because the search below starts *after* the heading and
    # a headless import has to start at the very first block.
    start = -1
    shape: _HeadingShape | None = None

    if heading is not None:
        wanted = _compile_heading(heading, path=path)
        for index, block in enumerate(blocks):
            if isinstance(block, WordParagraph) and wanted.fullmatch(block.text.strip()):
                start, shape = index, _shape_of(block)
                break
        else:
            raise DocParseError(f"{path}: no heading matching {heading!r} found")

    stop = None if end is None else _compile_heading(end, path=path)
    section: list[_Block] = []
    reached_end = False
    for block in blocks[start + 1 :]:
        if isinstance(block, WordParagraph) and _ends_section(block, stop, shape):
            reached_end = True
            break
        section.append(block)

    # A named end that never turns up is a mistake in the recipe, not a section
    # that happens to run to the last page: silently importing the rest of the
    # document would put another section's content under this heading.
    if stop is not None and not reached_end:
        raise DocParseError(f"{path}: no heading matching {end!r} follows {heading!r}")
    return section


def _compile_heading(text: str, *, path: Path) -> re.Pattern[str]:
    """Compile a ``begin``/``end`` value into the pattern it is matched as.

    Raises:
        DocParseError: if ``text`` is not a valid regular expression -- a recipe
            mistake, worth failing on before the document is even searched.
    """
    try:
        return re.compile(text.strip(), re.IGNORECASE)
    except re.error as exc:
        raise DocParseError(f"{path}: {text!r} is not a valid regular expression: {exc}") from exc


def _ends_section(
    paragraph: WordParagraph, stop: re.Pattern[str] | None, shape: _HeadingShape | None
) -> bool:
    if stop is not None:
        return stop.fullmatch(paragraph.text.strip()) is not None
    # No shape to compare against: the import began at the top of the document,
    # so there is no "another heading like this one" to look for.
    return shape is not None and _is_heading(paragraph, shape)


def _shape_of(paragraph: WordParagraph) -> _HeadingShape:
    first = next(iter(_runs(paragraph)), None)
    if first is None:
        return _HeadingShape(bold=False, size_pt=None)
    return _HeadingShape(bold=first.bold, size_pt=first.size_pt)


def _is_heading(paragraph: WordParagraph, shape: _HeadingShape) -> bool:
    """Whether ``paragraph`` is the next section heading, so the import stops.

    Deliberately narrow: same weight and same size as the heading we started
    from. Colour is not part of it -- in a real Document the headings differ there
    (one of them is the one Word autocorrected) without ceasing to be headings.
    """
    if _numbering_of(paragraph) is not None:
        return False
    first = next(iter(_runs(paragraph)), None)
    return first is not None and (first.bold, first.size_pt) == (shape.bold, shape.size_pt)


# -- conversion -----------------------------------------------------------


def _convert(block: _Block, numbering: _Numbering) -> RichBlock:
    if isinstance(block, WordTable):
        return _table(block, numbering)
    return _paragraph(block, numbering)


def _paragraph(paragraph: WordParagraph, numbering: _Numbering) -> RichParagraph:
    marker = _numbering_of(paragraph)
    level, ordered = (None, False) if marker is None else numbering.describe(*marker)
    return RichParagraph(runs=_merge(_runs(paragraph)), level=level, ordered=ordered)


def _table(table: WordTable, numbering: _Numbering) -> RichTable:
    return RichTable(
        rows=[
            [
                RichCell(
                    paragraphs=[
                        _paragraph(block, numbering)
                        for block in _iter_blocks(cell._tc, cell)
                        if isinstance(block, WordParagraph)
                    ]
                )
                for cell in row.cells
            ]
            for row in table.rows
        ],
        bordered=_is_bordered(table),
        column_widths=_column_widths(table),
    )


def _is_bordered(table: WordTable) -> bool:
    """Whether the table draws lines, from its own properties or its style."""
    elements = chain([table._tbl.tblPr], (style.element for style in _style_chain(table.style)))
    for element in elements:
        if element is None:
            continue
        properties = element if element.tag == qn("w:tblPr") else element.find(qn("w:tblPr"))
        borders = None if properties is None else properties.find(qn("w:tblBorders"))
        if borders is None:
            continue
        return any(edge.get(qn("w:val")) not in (None, "none", "nil") for edge in borders)
    return False


def _column_widths(table: WordTable) -> list[float]:
    widths = [column.width for column in table.columns]
    total = sum(width for width in widths if width is not None)
    if not total or any(width is None for width in widths):
        return []
    return [round(float(width) / float(total), 4) for width in widths if width is not None]


# -- runs and their formatting -------------------------------------------


def _runs(paragraph: WordParagraph) -> list[RichRun]:
    """The paragraph's runs, including those inside a hyperlink.

    ``Paragraph.runs`` skips a ``w:hyperlink``'s children, which would drop the
    link's text as well as the link.
    """
    collected: list[RichRun] = []
    for child in paragraph._p.iterchildren():
        if child.tag == qn("w:r"):
            collected.append(_run(Run(child, paragraph), paragraph, None))
        elif child.tag == qn("w:hyperlink"):
            url = _hyperlink_url(child, paragraph)
            collected.extend(
                _run(Run(element, paragraph), paragraph, url)
                for element in child.findall(qn("w:r"))
            )
    return [run for run in collected if run.text]


def _hyperlink_url(element: BaseOxmlElement, paragraph: WordParagraph) -> str | None:
    """The target of a ``w:hyperlink``, or None for an in-document anchor."""
    relationship_id = element.get(qn("r:id"))
    if relationship_id is None:
        return None
    relationship = paragraph.part.rels.get(relationship_id)
    return None if relationship is None else str(relationship.target_ref)


def _run(run: Run, paragraph: WordParagraph, link: str | None) -> RichRun:
    size = _resolve(run, paragraph, lambda font: font.size)
    color = _resolve(run, paragraph, lambda font: font.color.rgb)
    return RichRun(
        text=run.text,
        bold=bool(_resolve(run, paragraph, lambda font: font.bold)),
        italic=bool(_resolve(run, paragraph, lambda font: font.italic)),
        underline=bool(_resolve(run, paragraph, lambda font: font.underline)),
        strike=bool(_resolve(run, paragraph, lambda font: font.strike)),
        size_pt=None if size is None else float(size.pt),
        color=None if color is None else str(color),
        link=link,
    )


def _resolve(run: Run, paragraph: WordParagraph, pick: Callable[[Font], _T | None]) -> _T | None:
    """Read one font attribute the way Word does: direct, then the styles.

    ``None`` from python-docx means "not set here", which is exactly the cue to
    look one level up. An explicit ``False`` -- ``<w:b w:val="0"/>``, how Word
    switches bold *off* under a bold style -- is a value and stops the search.
    """
    direct = pick(run.font)
    if direct is not None:
        return direct
    for style in chain(_style_chain(run.style), _style_chain(paragraph.style)):
        inherited = pick(style.font)
        if inherited is not None:
            return inherited
    return None


def _style_chain(style: CharacterStyle | None) -> Iterator[CharacterStyle]:
    """A style and everything it is ``basedOn``, nearest first.

    ``CharacterStyle`` is the class that carries ``font`` and ``base_style``;
    paragraph and table styles both derive from it, so one walk serves all three.
    """
    seen: set[int] = set()
    current = style
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        yield current
        base = current.base_style
        current = base if isinstance(base, CharacterStyle) else None


def _merge(runs: Sequence[RichRun]) -> list[RichRun]:
    """Join neighbouring runs that carry identical formatting.

    Word splits runs at every spell-check marker and revision, so a single
    sentence can arrive as a dozen runs. Left alone they would become a dozen
    ``<span>``s or a dozen Word runs saying the same thing.
    """
    merged: list[RichRun] = []
    for run in runs:
        previous = merged[-1] if merged else None
        if previous is not None and _format_of(previous) == _format_of(run):
            merged[-1] = previous.model_copy(update={"text": previous.text + run.text})
        else:
            merged.append(run)
    return merged


def _format_of(run: RichRun) -> tuple[object, ...]:
    return (run.bold, run.italic, run.underline, run.strike, run.size_pt, run.color, run.link)


# -- numbering ------------------------------------------------------------


def _numbering_of(paragraph: WordParagraph) -> tuple[int, int] | None:
    """The paragraph's ``(num_id, level)``, or None if it is not a list item.

    ``numPr`` may sit on the paragraph or on any style it inherits from -- Word's
    own ``List Bullet`` style carries it -- so both have to be consulted.
    """
    elements = chain([paragraph._p], (style.element for style in _style_chain(paragraph.style)))
    for element in elements:
        properties = element.find(qn("w:pPr"))
        marker = None if properties is None else properties.find(qn("w:numPr"))
        if marker is None:
            continue
        num_id = _int_val(marker.find(qn("w:numId")))
        # numId 0 is how a paragraph opts *out* of the numbering its style gives it.
        if num_id is None or num_id == 0:
            return None
        return num_id, _int_val(marker.find(qn("w:ilvl"))) or 0
    return None


def _int_val(element: BaseOxmlElement | None) -> int | None:
    if element is None:
        return None
    value = element.get(qn("w:val"))
    try:
        return None if value is None else int(value)
    except ValueError:
        return None


class _Numbering:
    """The document's numbering definitions, read once per import."""

    def __init__(self, document: WordDocument) -> None:
        try:
            element = document.part.numbering_part.element
            self._element: BaseOxmlElement | None = element
        except (KeyError, NotImplementedError, ValueError):
            # A document that never used a list has no numbering part at all.
            self._element = None

    def describe(self, num_id: int, level: int) -> tuple[int, bool]:
        """Turn a ``(num_id, level)`` marker into ``(level, ordered)``."""
        return level, self._is_ordered(num_id, level)

    def _is_ordered(self, num_id: int, level: int) -> bool:
        definition = self._abstract(num_id)
        entry = None if definition is None else _child_where(definition, "w:lvl", "w:ilvl", level)
        format_element = None if entry is None else entry.find(qn("w:numFmt"))
        value = None if format_element is None else format_element.get(qn("w:val"))
        return value not in (None, "bullet", "none")

    def _abstract(self, num_id: int) -> BaseOxmlElement | None:
        """The ``w:abstractNum`` the given ``w:num`` points at."""
        if self._element is None:
            return None
        num = _child_where(self._element, "w:num", "w:numId", num_id)
        reference = None if num is None else num.find(qn("w:abstractNumId"))
        abstract_id = None if reference is None else reference.get(qn("w:val"))
        if abstract_id is None:
            return None
        return _child_where(self._element, "w:abstractNum", "w:abstractNumId", abstract_id)


def _child_where(
    parent: BaseOxmlElement, tag: str, attribute: str, value: object
) -> BaseOxmlElement | None:
    """The first ``tag`` child of ``parent`` whose ``attribute`` equals ``value``."""
    for child in parent.findall(qn(tag)):
        element: BaseOxmlElement = child
        if element.get(qn(attribute)) == str(value):
            return element
    return None


__all__ = ["LOCK_PREFIX", "load_section"]
